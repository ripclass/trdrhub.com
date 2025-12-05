"""
Certificate Generators

Specialized generators for:
- GSP Form A (Generalized System of Preferences)
- EUR.1 Movement Certificate
- Other preferential origin certificates
"""

import io
import logging
from datetime import date, datetime
from decimal import Decimal
from typing import Dict, Any, Optional, List

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer,
    PageBreak, Image, HRFlowable
)
from reportlab.pdfgen import canvas

logger = logging.getLogger(__name__)


class GSPFormAGenerator:
    """
    Generator for GSP Form A - Certificate of Origin.
    
    Used for preferential tariff treatment under GSP schemes
    for exports to developed countries.
    """
    
    def __init__(self):
        self.page_width, self.page_height = A4
        self.margin = 15 * mm
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for GSP Form A"""
        self.styles.add(ParagraphStyle(
            name='GSPHeader',
            fontName='Helvetica-Bold',
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=6,
        ))
        self.styles.add(ParagraphStyle(
            name='GSPBoxLabel',
            fontName='Helvetica',
            fontSize=7,
            textColor=colors.black,
        ))
        self.styles.add(ParagraphStyle(
            name='GSPBoxContent',
            fontName='Helvetica',
            fontSize=9,
            leading=11,
        ))
        self.styles.add(ParagraphStyle(
            name='GSPSmall',
            fontName='Helvetica',
            fontSize=6,
            textColor=colors.gray,
        ))
    
    def generate(self, data: Dict[str, Any]) -> bytes:
        """
        Generate GSP Form A PDF.
        
        Required data:
            - exporter_name, exporter_address
            - consignee_name, consignee_address
            - country_of_origin
            - destination_country
            - goods_description
            - hs_code
            - quantity
            - gross_weight
            - invoice_number, invoice_date
            - transport_details (vessel, route)
            - origin_criterion
        """
        buffer = io.BytesIO()
        
        # Create canvas for custom drawing
        c = canvas.Canvas(buffer, pagesize=A4)
        
        # Draw green border (standard for GSP Form A)
        c.setStrokeColor(colors.HexColor("#006400"))
        c.setLineWidth(2)
        c.rect(10*mm, 10*mm, self.page_width - 20*mm, self.page_height - 20*mm)
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(self.page_width/2, self.page_height - 25*mm, 
                          "GENERALIZED SYSTEM OF PREFERENCES")
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(self.page_width/2, self.page_height - 32*mm, 
                          "CERTIFICATE OF ORIGIN")
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(self.page_width/2, self.page_height - 39*mm, 
                          "(Combined declaration and certificate)")
        c.setFont("Helvetica-Bold", 18)
        c.drawCentredString(self.page_width/2, self.page_height - 48*mm, "FORM A")
        
        # Reference number
        c.setFont("Helvetica", 9)
        c.drawRightString(self.page_width - 15*mm, self.page_height - 20*mm,
                         f"Reference No: {data.get('reference_number', '')}")
        
        # Draw boxes
        y_start = self.page_height - 60*mm
        box_height = 30*mm
        col_width = (self.page_width - 30*mm) / 2
        
        # Box 1: Exporter
        self._draw_box(c, 15*mm, y_start, col_width, box_height, "1")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Goods consigned from (Exporter's business name, address, country)")
        c.setFont("Helvetica", 9)
        self._draw_multiline(c, 20*mm, y_start - 12*mm, 
                            f"{data.get('exporter_name', '')}\n{data.get('exporter_address', '')}")
        
        # Box 2: Reference/For official use
        self._draw_box(c, 15*mm + col_width, y_start, col_width, box_height, "2")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - 5*mm, "Reference No. (For official use)")
        
        # Box 3: Consignee
        y_start -= box_height
        self._draw_box(c, 15*mm, y_start, col_width, box_height, "3")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Goods consigned to (Consignee's name, address, country)")
        c.setFont("Helvetica", 9)
        self._draw_multiline(c, 20*mm, y_start - 12*mm,
                            f"{data.get('consignee_name', '')}\n{data.get('consignee_address', '')}")
        
        # Box 4: Country of Origin
        self._draw_box(c, 15*mm + col_width, y_start, col_width, box_height / 2, "4")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - 5*mm, "Country of Origin")
        c.setFont("Helvetica-Bold", 10)
        c.drawString(20*mm + col_width, y_start - 15*mm, data.get('country_of_origin', ''))
        
        # Box 5: Destination Country
        self._draw_box(c, 15*mm + col_width, y_start - box_height/2, col_width, box_height / 2, "5")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - box_height/2 - 5*mm, "For official use")
        
        # Box 6: Transport details
        y_start -= box_height
        self._draw_box(c, 15*mm, y_start, self.page_width - 30*mm, 20*mm, "6")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Means of transport and route (as far as known)")
        c.setFont("Helvetica", 9)
        c.drawString(20*mm, y_start - 12*mm, data.get('transport_details', ''))
        
        # Goods table (Boxes 7-11)
        y_start -= 25*mm
        table_height = 80*mm
        
        # Headers
        headers = [
            ("7", "Item\nnumber", 15*mm),
            ("8", "Marks and numbers\nof packages", 35*mm),
            ("9", "Number and kind of packages;\ndescription of goods", 60*mm),
            ("10", "Origin\ncriterion", 20*mm),
            ("11", "Gross weight\nor quantity", 25*mm),
            ("12", "Number and\ndate of invoice", 30*mm),
        ]
        
        x_pos = 15*mm
        for box_num, header, width in headers:
            c.setStrokeColor(colors.HexColor("#006400"))
            c.rect(x_pos, y_start - table_height, width, table_height)
            c.setFont("Helvetica", 7)
            c.drawCentredString(x_pos + width/2, y_start - 5*mm, box_num)
            c.setFont("Helvetica", 6)
            for i, line in enumerate(header.split('\n')):
                c.drawCentredString(x_pos + width/2, y_start - 12*mm - i*7, line)
            x_pos += width
        
        # Goods content
        c.setFont("Helvetica", 9)
        c.drawString(20*mm, y_start - 35*mm, "1")  # Item number
        c.drawString(35*mm, y_start - 35*mm, data.get('shipping_marks', 'N/M'))
        
        # Description (with HS code)
        desc = f"{data.get('goods_description', '')}\n\nHS Code: {data.get('hs_code', '')}"
        self._draw_multiline(c, 70*mm, y_start - 35*mm, desc, max_width=55*mm)
        
        c.drawCentredString(140*mm, y_start - 35*mm, data.get('origin_criterion', 'P'))
        c.drawString(150*mm, y_start - 35*mm, f"{data.get('gross_weight', '')} KGS")
        c.drawString(175*mm, y_start - 35*mm, f"{data.get('invoice_number', '')}\n{data.get('invoice_date', '')}")
        
        # Declaration box (Box 12)
        y_start -= table_height + 5*mm
        decl_height = 35*mm
        self._draw_box(c, 15*mm, y_start, col_width, decl_height, "12")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Declaration by the exporter")
        c.setFont("Helvetica", 8)
        declaration = """The undersigned hereby declares that the above
details and statements are correct; that all the
goods were produced in"""
        self._draw_multiline(c, 20*mm, y_start - 12*mm, declaration)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(20*mm, y_start - 28*mm, data.get('country_of_origin', ''))
        c.setFont("Helvetica", 8)
        c.drawString(55*mm, y_start - 28*mm, f"(Country)")
        
        # Certification box (Box 13)
        self._draw_box(c, 15*mm + col_width, y_start, col_width, decl_height, "13")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - 5*mm, "Certification")
        c.setFont("Helvetica", 8)
        certification = """It is hereby certified, on the basis of control
carried out, that the declaration by the exporter
is correct."""
        self._draw_multiline(c, 20*mm + col_width, y_start - 12*mm, certification)
        
        # Signature boxes
        y_start -= decl_height + 5*mm
        sig_height = 25*mm
        
        # Exporter signature
        self._draw_box(c, 15*mm, y_start, col_width, sig_height, "")
        c.setFont("Helvetica", 8)
        c.drawString(20*mm, y_start - 8*mm, f"Place and date: {data.get('place', '')}, {data.get('date', datetime.now().strftime('%Y-%m-%d'))}")
        c.drawString(20*mm, y_start - 20*mm, "Signature of exporter: _____________________")
        
        # Official stamp
        self._draw_box(c, 15*mm + col_width, y_start, col_width, sig_height, "")
        c.setFont("Helvetica", 8)
        c.drawString(20*mm + col_width, y_start - 8*mm, f"Place and date:")
        c.drawString(20*mm + col_width, y_start - 15*mm, "Signature and stamp of")
        c.drawString(20*mm + col_width, y_start - 20*mm, "certifying authority: _____________________")
        
        c.save()
        buffer.seek(0)
        return buffer.getvalue()
    
    def _draw_box(self, c, x, y, width, height, box_number: str):
        """Draw a numbered box"""
        c.setStrokeColor(colors.HexColor("#006400"))
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)
        if box_number:
            c.setFont("Helvetica", 6)
            c.setFillColor(colors.gray)
            c.drawString(x + 1*mm, y - 3*mm, box_number)
            c.setFillColor(colors.black)
    
    def _draw_multiline(self, c, x, y, text, max_width=None, line_height=10):
        """Draw multiline text"""
        lines = text.split('\n')
        for i, line in enumerate(lines[:6]):  # Limit to 6 lines
            c.drawString(x, y - i*line_height, line[:60])  # Limit line length


class EUR1CertificateGenerator:
    """
    Generator for EUR.1 Movement Certificate.
    
    Used for preferential origin in trade between EU and
    partner countries with free trade agreements.
    """
    
    def __init__(self):
        self.page_width, self.page_height = A4
        self.margin = 15 * mm
        self.styles = getSampleStyleSheet()
    
    def generate(self, data: Dict[str, Any]) -> bytes:
        """
        Generate EUR.1 Movement Certificate PDF.
        
        Required data:
            - exporter_name, exporter_address, exporter_country
            - consignee_name, consignee_address, consignee_country
            - goods_description
            - gross_weight
            - invoice_details
            - origin_country
            - destination_country
        """
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        
        # Title with EU styling (blue and gold)
        c.setFillColor(colors.HexColor("#003399"))
        c.setFont("Helvetica-Bold", 18)
        c.drawCentredString(self.page_width/2, self.page_height - 25*mm, 
                          "MOVEMENT CERTIFICATE")
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(self.page_width/2, self.page_height - 35*mm, "EUR.1")
        
        # Certificate number
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 9)
        c.drawRightString(self.page_width - 15*mm, self.page_height - 15*mm,
                         f"No: {data.get('certificate_number', '')}")
        
        # Draw the standard EUR.1 boxes
        y_start = self.page_height - 50*mm
        col_width = (self.page_width - 30*mm) / 2
        
        # Box 1: Exporter
        self._draw_labeled_box(c, 15*mm, y_start, col_width, 35*mm,
            "1. Exporter (Name, full address, country)",
            f"{data.get('exporter_name', '')}\n{data.get('exporter_address', '')}")
        
        # Box 2: Certificate number (for customs)
        self._draw_labeled_box(c, 15*mm + col_width, y_start, col_width, 35*mm,
            "2. Certificate used in preferential trade between",
            f"{data.get('origin_country', '')} and {data.get('destination_country', '')}")
        
        y_start -= 35*mm
        
        # Box 3: Consignee
        self._draw_labeled_box(c, 15*mm, y_start, col_width, 35*mm,
            "3. Consignee (Name, full address, country)",
            f"{data.get('consignee_name', '')}\n{data.get('consignee_address', '')}")
        
        # Box 4: Country of origin
        self._draw_labeled_box(c, 15*mm + col_width, y_start, col_width, 17*mm,
            "4. Country, group of countries or territory in which the products are considered as originating",
            data.get('origin_country', ''))
        
        # Box 5: Destination country
        self._draw_labeled_box(c, 15*mm + col_width, y_start - 18*mm, col_width, 17*mm,
            "5. Country, group of countries or territory of destination",
            data.get('destination_country', ''))
        
        y_start -= 40*mm
        
        # Box 6: Transport details
        self._draw_labeled_box(c, 15*mm, y_start, self.page_width - 30*mm, 20*mm,
            "6. Transport details (Optional)",
            data.get('transport_details', ''))
        
        y_start -= 25*mm
        
        # Box 7: Remarks
        self._draw_labeled_box(c, 15*mm, y_start, self.page_width - 30*mm, 15*mm,
            "7. Remarks",
            data.get('remarks', ''))
        
        y_start -= 20*mm
        
        # Box 8: Goods description (large box)
        goods_box_height = 60*mm
        self._draw_labeled_box(c, 15*mm, y_start, self.page_width - 30*mm, goods_box_height,
            "8. Item number; Marks and numbers; Number and kind of packages; Description of goods",
            f"Marks: {data.get('shipping_marks', 'N/M')}\n\n{data.get('goods_description', '')}")
        
        # Add columns for gross weight and invoice
        c.setStrokeColor(colors.black)
        c.line(140*mm, y_start, 140*mm, y_start - goods_box_height)
        c.line(165*mm, y_start, 165*mm, y_start - goods_box_height)
        
        c.setFont("Helvetica", 7)
        c.drawCentredString(152*mm, y_start - 5*mm, "9. Gross weight (kg)")
        c.drawCentredString(180*mm, y_start - 5*mm, "10. Invoices")
        
        c.setFont("Helvetica", 9)
        c.drawCentredString(152*mm, y_start - 30*mm, f"{data.get('gross_weight', '')}")
        c.drawCentredString(180*mm, y_start - 30*mm, f"{data.get('invoice_number', '')}")
        
        y_start -= goods_box_height + 5*mm
        
        # Declaration and certification boxes
        decl_height = 35*mm
        
        # Box 11: Customs endorsement
        self._draw_labeled_box(c, 15*mm, y_start, col_width, decl_height,
            "11. CUSTOMS ENDORSEMENT\nDeclaration certified\nExport document:\nType: .............. No: ..............\nCustoms office: ...........................\nIssuing country: ...........................",
            "")
        
        # Add stamp placeholder
        c.setStrokeColor(colors.gray)
        c.setDash(3, 3)
        c.circle(15*mm + col_width/2, y_start - decl_height + 15*mm, 12*mm)
        c.setDash()
        c.setFont("Helvetica", 6)
        c.drawCentredString(15*mm + col_width/2, y_start - decl_height + 5*mm, "Stamp")
        
        # Box 12: Declaration by exporter
        self._draw_labeled_box(c, 15*mm + col_width, y_start, col_width, decl_height,
            "12. DECLARATION BY THE EXPORTER",
            f"""I, the undersigned, declare that the goods 
described above meet the conditions required
for the issue of this certificate.

Place and date: {data.get('place', '')}, {data.get('date', datetime.now().strftime('%Y-%m-%d'))}

Signature: _____________________""")
        
        c.save()
        buffer.seek(0)
        return buffer.getvalue()
    
    def _draw_labeled_box(self, c, x, y, width, height, label: str, content: str):
        """Draw a box with label and content"""
        c.setStrokeColor(colors.black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)
        
        # Label in small text
        c.setFont("Helvetica", 7)
        label_lines = label.split('\n')
        for i, line in enumerate(label_lines[:4]):
            c.drawString(x + 2*mm, y - 4*mm - i*8, line)
        
        # Content
        c.setFont("Helvetica", 9)
        content_lines = content.split('\n')
        content_y = y - 4*mm - len(label_lines[:4])*8 - 5*mm
        for i, line in enumerate(content_lines[:6]):
            c.drawString(x + 2*mm, content_y - i*10, line[:50])


class RCEPCertificateGenerator:
    """
    Generator for RCEP Certificate of Origin.
    
    Regional Comprehensive Economic Partnership certificate
    for trade among ASEAN+5 countries (China, Japan, Korea, Australia, NZ).
    """
    
    def __init__(self):
        self.page_width, self.page_height = A4
        self.margin = 15 * mm
    
    def generate(self, data: Dict[str, Any]) -> bytes:
        """
        Generate RCEP Certificate of Origin PDF.
        
        Required data:
            - exporter_name, exporter_address
            - producer_name, producer_address (if different from exporter)
            - importer_name, importer_address
            - goods_description, hs_code
            - origin_country
            - destination_country
            - invoice_number, invoice_date
            - rcep_origin_criterion
        """
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        
        # Blue header band (RCEP branding)
        c.setFillColor(colors.HexColor("#1a5276"))
        c.rect(0, self.page_height - 50*mm, self.page_width, 50*mm, fill=True, stroke=False)
        
        # Title
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(self.page_width/2, self.page_height - 20*mm, 
                          "REGIONAL COMPREHENSIVE ECONOMIC PARTNERSHIP")
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(self.page_width/2, self.page_height - 30*mm, 
                          "CERTIFICATE OF ORIGIN")
        c.setFont("Helvetica", 10)
        c.drawCentredString(self.page_width/2, self.page_height - 40*mm, 
                          "(Combined Declaration and Certificate) - FORM RCEP")
        
        # Certificate number
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 9)
        ref_no = data.get('reference_number', f"RCEP-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}")
        c.drawRightString(self.page_width - 15*mm, self.page_height - 55*mm, f"Reference No: {ref_no}")
        
        y_start = self.page_height - 65*mm
        col_width = (self.page_width - 30*mm) / 2
        box_height = 30*mm
        
        # Box 1: Exporter
        self._draw_box(c, 15*mm, y_start, col_width, box_height, "1")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Exporter's Name, Address and Country")
        c.setFont("Helvetica", 9)
        exporter = f"{data.get('exporter_name', '')}\n{data.get('exporter_address', '')}"
        self._draw_multiline(c, 20*mm, y_start - 12*mm, exporter)
        
        # Box 2: Producer (if different)
        self._draw_box(c, 15*mm + col_width, y_start, col_width, box_height, "2")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - 5*mm, "Producer's Name and Address (if known)")
        c.setFont("Helvetica", 9)
        producer = data.get('producer_name', 'SAME AS EXPORTER')
        if data.get('producer_address'):
            producer += f"\n{data.get('producer_address', '')}"
        self._draw_multiline(c, 20*mm + col_width, y_start - 12*mm, producer)
        
        y_start -= box_height
        
        # Box 3: Importer
        self._draw_box(c, 15*mm, y_start, col_width, box_height, "3")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Importer's/Consignee's Name, Address and Country")
        c.setFont("Helvetica", 9)
        importer = f"{data.get('importer_name', data.get('consignee_name', ''))}\n{data.get('importer_address', data.get('consignee_address', ''))}"
        self._draw_multiline(c, 20*mm, y_start - 12*mm, importer)
        
        # Box 4: Importing Party
        self._draw_box(c, 15*mm + col_width, y_start, col_width, box_height/2, "4")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - 5*mm, "Importing Party")
        c.setFont("Helvetica-Bold", 10)
        c.drawString(20*mm + col_width, y_start - 15*mm, data.get('destination_country', ''))
        
        # Box 5: Exporting Party
        self._draw_box(c, 15*mm + col_width, y_start - box_height/2, col_width, box_height/2, "5")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - box_height/2 - 5*mm, "Exporting Party")
        c.setFont("Helvetica-Bold", 10)
        c.drawString(20*mm + col_width, y_start - box_height/2 - 15*mm, data.get('origin_country', ''))
        
        y_start -= box_height
        
        # Box 6-10: Goods table
        table_height = 70*mm
        c.setStrokeColor(colors.black)
        c.rect(15*mm, y_start - table_height, self.page_width - 30*mm, table_height)
        
        # Table headers
        headers = [
            ("6", "Item\nNo.", 15*mm),
            ("7", "HS Code\n(6-digit)", 35*mm),
            ("8", "Description of goods", 120*mm),
            ("9", "RCEP\nCriterion", 30*mm),
            ("10", "Gross/Net\nWeight", 35*mm),
            ("11", "Invoice No.\n& Date", 50*mm),
        ]
        
        x_pos = 15*mm
        for box_num, header, width in headers:
            c.setStrokeColor(colors.black)
            c.rect(x_pos, y_start - table_height, width, table_height)
            c.setFont("Helvetica", 6)
            c.drawCentredString(x_pos + width/2, y_start - 4*mm, box_num)
            c.setFont("Helvetica", 7)
            for i, line in enumerate(header.split('\n')):
                c.drawCentredString(x_pos + width/2, y_start - 10*mm - i*7, line)
            x_pos += width
        
        # Content row
        c.setFont("Helvetica", 9)
        content_y = y_start - 35*mm
        c.drawCentredString(15*mm + 7.5*mm, content_y, "1")  # Item number
        c.drawCentredString(15*mm + 15*mm + 17.5*mm, content_y, data.get('hs_code', '')[:6])  # HS Code
        
        # Description (wrap text)
        desc = data.get('goods_description', '')[:100]
        c.drawString(15*mm + 50*mm + 5*mm, content_y, desc[:50])
        if len(desc) > 50:
            c.drawString(15*mm + 50*mm + 5*mm, content_y - 10, desc[50:100])
        
        # RCEP Criterion
        criterion = data.get('rcep_origin_criterion', data.get('origin_criterion', 'WO'))
        c.drawCentredString(15*mm + 170*mm + 15*mm, content_y, criterion)
        
        # Weight
        weight = f"{data.get('gross_weight', '')} KG"
        c.drawCentredString(15*mm + 200*mm + 17.5*mm, content_y, weight)
        
        # Invoice
        inv = f"{data.get('invoice_number', '')}\n{data.get('invoice_date', '')}"
        c.drawCentredString(15*mm + 235*mm + 25*mm, content_y, data.get('invoice_number', ''))
        c.setFont("Helvetica", 8)
        c.drawCentredString(15*mm + 235*mm + 25*mm, content_y - 10, str(data.get('invoice_date', '')))
        
        y_start -= table_height + 5*mm
        
        # Box 12: Declaration
        decl_height = 40*mm
        self._draw_box(c, 15*mm, y_start, col_width, decl_height, "12")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm, y_start - 5*mm, "Declaration by the exporter or producer")
        
        declaration = """The undersigned hereby declares that the above 
details and statements are correct, that all the 
goods were produced in

and that they comply with the origin requirements 
specified for those goods in RCEP."""
        c.setFont("Helvetica", 8)
        for i, line in enumerate(declaration.strip().split('\n')):
            c.drawString(20*mm, y_start - 12*mm - i*9, line)
        
        c.setFont("Helvetica-Bold", 9)
        c.drawString(20*mm, y_start - 30*mm, data.get('origin_country', ''))
        
        # Box 13: Certification
        self._draw_box(c, 15*mm + col_width, y_start, col_width, decl_height, "13")
        c.setFont("Helvetica", 7)
        c.drawString(20*mm + col_width, y_start - 5*mm, "Certification")
        
        certification = """It is hereby certified, on the basis of control 
carried out, that the declaration by the exporter 
or producer is correct."""
        c.setFont("Helvetica", 8)
        for i, line in enumerate(certification.strip().split('\n')):
            c.drawString(20*mm + col_width, y_start - 12*mm - i*9, line)
        
        y_start -= decl_height + 5*mm
        
        # Signature boxes
        sig_height = 25*mm
        
        # Exporter signature
        self._draw_box(c, 15*mm, y_start, col_width, sig_height, "")
        c.setFont("Helvetica", 8)
        c.drawString(20*mm, y_start - 8*mm, f"Place: {data.get('place', '')}")
        c.drawString(20*mm, y_start - 15*mm, f"Date: {data.get('date', datetime.now().strftime('%Y-%m-%d'))}")
        c.drawString(20*mm, y_start - 22*mm, "Signature: _____________________")
        
        # Issuing authority
        self._draw_box(c, 15*mm + col_width, y_start, col_width, sig_height, "")
        c.setFont("Helvetica", 8)
        c.drawString(20*mm + col_width, y_start - 8*mm, "Issuing Authority:")
        c.drawString(20*mm + col_width, y_start - 15*mm, "Date: _____________________")
        c.drawString(20*mm + col_width, y_start - 22*mm, "Signature & Stamp: _____________")
        
        # RCEP member countries footer
        c.setFont("Helvetica", 6)
        c.setFillColor(colors.grey)
        members = "RCEP Members: Australia, Brunei, Cambodia, China, Indonesia, Japan, Korea, Laos, Malaysia, Myanmar, New Zealand, Philippines, Singapore, Thailand, Vietnam"
        c.drawCentredString(self.page_width/2, 20*mm, members)
        
        c.save()
        buffer.seek(0)
        return buffer.getvalue()
    
    def _draw_box(self, c, x, y, width, height, box_number: str):
        """Draw a numbered box"""
        c.setStrokeColor(colors.black)
        c.setLineWidth(0.5)
        c.rect(x, y - height, width, height)
        if box_number:
            c.setFont("Helvetica", 6)
            c.setFillColor(colors.grey)
            c.drawString(x + 1*mm, y - 3*mm, box_number)
            c.setFillColor(colors.black)
    
    def _draw_multiline(self, c, x, y, text, max_width=None, line_height=10):
        """Draw multiline text"""
        lines = text.split('\n')
        for i, line in enumerate(lines[:5]):
            c.drawString(x, y - i*line_height, line[:55])


class DocumentExportService:
    """
    Service for exporting documents to Word/Excel formats.
    """
    
    def export_to_docx(self, data: Dict[str, Any], template_name: str) -> bytes:
        """
        Export document data to Word format.
        
        Uses python-docx library.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx required for Word export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading(template_name.replace("_", " ").upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parties
        doc.add_heading("Parties", level=1)
        doc.add_paragraph(f"Beneficiary: {data.get('beneficiary_name', '')}")
        doc.add_paragraph(f"Address: {data.get('beneficiary_address', '')}")
        doc.add_paragraph()
        doc.add_paragraph(f"Applicant: {data.get('applicant_name', '')}")
        doc.add_paragraph(f"Address: {data.get('applicant_address', '')}")
        
        # Shipment Details
        doc.add_heading("Shipment Details", level=1)
        doc.add_paragraph(f"Port of Loading: {data.get('port_of_loading', '')}")
        doc.add_paragraph(f"Port of Discharge: {data.get('port_of_discharge', '')}")
        doc.add_paragraph(f"Vessel: {data.get('vessel_name', '')}")
        doc.add_paragraph(f"B/L Number: {data.get('bl_number', '')}")
        
        # Goods (Table)
        if 'line_items' in data and data['line_items']:
            doc.add_heading("Goods", level=1)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Headers
            headers = table.rows[0].cells
            headers[0].text = "Description"
            headers[1].text = "Qty"
            headers[2].text = "Unit"
            headers[3].text = "Unit Price"
            headers[4].text = "Total"
            
            # Data rows
            for item in data['line_items']:
                row = table.add_row().cells
                row[0].text = str(item.get('description', ''))
                row[1].text = str(item.get('quantity', ''))
                row[2].text = str(item.get('unit', ''))
                row[3].text = str(item.get('unit_price', ''))
                row[4].text = str(item.get('total_price', ''))
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_xlsx(self, data: Dict[str, Any], template_name: str) -> bytes:
        """
        Export document data to Excel format.
        
        Uses openpyxl library.
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
        except ImportError:
            raise ImportError("openpyxl required for Excel export. Install with: pip install openpyxl")
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Document Data"
        
        # Header
        ws['A1'] = template_name.replace("_", " ").upper()
        ws['A1'].font = Font(bold=True, size=14)
        
        # Parties
        ws['A3'] = "PARTIES"
        ws['A3'].font = Font(bold=True)
        ws['A4'] = "Beneficiary"
        ws['B4'] = data.get('beneficiary_name', '')
        ws['A5'] = "Beneficiary Address"
        ws['B5'] = data.get('beneficiary_address', '')
        ws['A6'] = "Applicant"
        ws['B6'] = data.get('applicant_name', '')
        ws['A7'] = "Applicant Address"
        ws['B7'] = data.get('applicant_address', '')
        
        # Shipment
        ws['A9'] = "SHIPMENT DETAILS"
        ws['A9'].font = Font(bold=True)
        ws['A10'] = "Port of Loading"
        ws['B10'] = data.get('port_of_loading', '')
        ws['A11'] = "Port of Discharge"
        ws['B11'] = data.get('port_of_discharge', '')
        ws['A12'] = "Vessel"
        ws['B12'] = data.get('vessel_name', '')
        ws['A13'] = "B/L Number"
        ws['B13'] = data.get('bl_number', '')
        
        # Line items
        if 'line_items' in data and data['line_items']:
            ws['A15'] = "LINE ITEMS"
            ws['A15'].font = Font(bold=True)
            
            # Headers
            headers = ['Description', 'Quantity', 'Unit', 'Unit Price', 'Total']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=16, column=col, value=header)
                cell.font = Font(bold=True)
            
            # Data
            for row_num, item in enumerate(data['line_items'], 17):
                ws.cell(row=row_num, column=1, value=item.get('description', ''))
                ws.cell(row=row_num, column=2, value=item.get('quantity', ''))
                ws.cell(row=row_num, column=3, value=item.get('unit', ''))
                ws.cell(row=row_num, column=4, value=item.get('unit_price', ''))
                ws.cell(row=row_num, column=5, value=item.get('total_price', ''))
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 40
        
        # Save
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# Singleton instances
_gsp_generator: Optional[GSPFormAGenerator] = None
_eur1_generator: Optional[EUR1CertificateGenerator] = None
_rcep_generator: Optional[RCEPCertificateGenerator] = None
_export_service: Optional[DocumentExportService] = None


def get_gsp_generator() -> GSPFormAGenerator:
    global _gsp_generator
    if _gsp_generator is None:
        _gsp_generator = GSPFormAGenerator()
    return _gsp_generator


def get_eur1_generator() -> EUR1CertificateGenerator:
    global _eur1_generator
    if _eur1_generator is None:
        _eur1_generator = EUR1CertificateGenerator()
    return _eur1_generator


def get_rcep_generator() -> RCEPCertificateGenerator:
    global _rcep_generator
    if _rcep_generator is None:
        _rcep_generator = RCEPCertificateGenerator()
    return _rcep_generator


def get_export_service() -> DocumentExportService:
    global _export_service
    if _export_service is None:
        _export_service = DocumentExportService()
    return _export_service

