"""
LC Builder API Router

Endpoints for creating and managing LC applications.
"""

import uuid
import logging
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, HTTPException, Depends, Query, BackgroundTasks
from sqlalchemy.orm import Session
from pydantic import BaseModel, Field

from app.database import get_db
from app.models.user import User
from app.models.lc_builder import (
    LCApplication, LCDocumentRequirement, LCApplicationVersion,
    LCClause, LCTemplate, ApplicantProfile, BeneficiaryProfile,
    LCType, LCStatus, PaymentTerms, ConfirmationInstructions
)
from app.routers.auth import get_current_user
from app.services.lc_clause_library import (
    LCClauseLibrary, ClauseCategory, RiskLevel, BiasIndicator
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/lc-builder", tags=["lc-builder"])


# ============================================================================
# Pydantic Models
# ============================================================================

class PartyInfo(BaseModel):
    name: str
    address: Optional[str] = None
    country: Optional[str] = None
    contact: Optional[str] = None


class BankInfo(BaseModel):
    name: Optional[str] = None
    swift: Optional[str] = None


class DocumentRequirementCreate(BaseModel):
    document_type: str
    description: Optional[str] = None
    copies_original: int = 1
    copies_copy: int = 0
    specific_requirements: Optional[str] = None
    is_required: bool = True


class LCApplicationCreate(BaseModel):
    """Create a new LC application"""
    name: Optional[str] = None
    lc_type: str = "documentary"
    
    # Amount
    currency: str = "USD"
    amount: float
    tolerance_plus: float = 0
    tolerance_minus: float = 0
    
    # Parties
    applicant: PartyInfo
    beneficiary: PartyInfo
    issuing_bank: Optional[BankInfo] = None
    advising_bank: Optional[BankInfo] = None
    
    # Shipment
    port_of_loading: Optional[str] = None
    port_of_discharge: Optional[str] = None
    place_of_delivery: Optional[str] = None
    latest_shipment_date: Optional[datetime] = None
    incoterms: Optional[str] = None
    incoterms_place: Optional[str] = None
    partial_shipments: bool = True
    transhipment: bool = True
    
    # Goods
    goods_description: str
    hs_code: Optional[str] = None
    quantity: Optional[str] = None
    unit_price: Optional[str] = None
    
    # Payment
    payment_terms: str = "sight"
    usance_days: Optional[int] = None
    usance_from: Optional[str] = None
    
    # Validity
    expiry_date: datetime
    expiry_place: Optional[str] = None
    presentation_period: int = 21
    confirmation_instructions: str = "without"
    
    # Documents & Clauses
    documents_required: List[DocumentRequirementCreate] = []
    additional_conditions: List[str] = []
    selected_clause_ids: List[str] = []
    
    # Template
    template_id: Optional[str] = None


class LCApplicationUpdate(BaseModel):
    """Update LC application"""
    name: Optional[str] = None
    lc_type: Optional[str] = None
    status: Optional[str] = None
    
    # Amount
    currency: Optional[str] = None
    amount: Optional[float] = None
    tolerance_plus: Optional[float] = None
    tolerance_minus: Optional[float] = None
    
    # Parties
    applicant_name: Optional[str] = None
    applicant_address: Optional[str] = None
    applicant_country: Optional[str] = None
    beneficiary_name: Optional[str] = None
    beneficiary_address: Optional[str] = None
    beneficiary_country: Optional[str] = None
    
    # Banks
    issuing_bank_name: Optional[str] = None
    issuing_bank_swift: Optional[str] = None
    advising_bank_name: Optional[str] = None
    advising_bank_swift: Optional[str] = None
    
    # Shipment
    port_of_loading: Optional[str] = None
    port_of_discharge: Optional[str] = None
    latest_shipment_date: Optional[datetime] = None
    incoterms: Optional[str] = None
    incoterms_place: Optional[str] = None
    partial_shipments: Optional[bool] = None
    transhipment: Optional[bool] = None
    
    # Goods
    goods_description: Optional[str] = None
    hs_code: Optional[str] = None
    
    # Payment
    payment_terms: Optional[str] = None
    usance_days: Optional[int] = None
    
    # Validity
    expiry_date: Optional[datetime] = None
    expiry_place: Optional[str] = None
    presentation_period: Optional[int] = None
    confirmation_instructions: Optional[str] = None
    
    # Documents & Clauses
    documents_required: Optional[List[DocumentRequirementCreate]] = None
    additional_conditions: Optional[List[str]] = None
    selected_clause_ids: Optional[List[str]] = None


class ValidationIssue(BaseModel):
    field: str
    severity: str  # error, warning, info
    message: str
    suggestion: Optional[str] = None


class ValidationResult(BaseModel):
    is_valid: bool
    issues: List[ValidationIssue]
    risk_score: float
    risk_details: Dict[str, Any]


class ClauseResponse(BaseModel):
    code: str
    category: str
    subcategory: str
    title: str
    clause_text: str
    plain_english: str
    risk_level: str
    bias: str
    risk_notes: str
    bank_acceptance: float
    tags: List[str]


# ============================================================================
# Helper Functions
# ============================================================================

def generate_reference_number() -> str:
    """Generate unique LC reference number"""
    timestamp = datetime.now().strftime("%Y%m%d")
    unique_id = str(uuid.uuid4())[:8].upper()
    return f"LC-{timestamp}-{unique_id}"


def validate_lc_application(lc: LCApplication) -> ValidationResult:
    """Validate LC application and calculate risk score"""
    issues = []
    risk_factors = []
    
    # Date validations
    if lc.latest_shipment_date and lc.expiry_date:
        days_between = (lc.expiry_date - lc.latest_shipment_date).days
        if days_between < 21:
            issues.append(ValidationIssue(
                field="expiry_date",
                severity="warning",
                message=f"Only {days_between} days between shipment and expiry (minimum 21 recommended)",
                suggestion="Extend expiry date to allow time for document preparation and presentation"
            ))
            risk_factors.append(("tight_timeline", 15))
        
        if days_between < 0:
            issues.append(ValidationIssue(
                field="expiry_date",
                severity="error",
                message="Expiry date is before shipment date",
                suggestion="Expiry must be after latest shipment date"
            ))
    
    # Amount validations
    if lc.amount <= 0:
        issues.append(ValidationIssue(
            field="amount",
            severity="error",
            message="Amount must be greater than zero"
        ))
    
    if lc.tolerance_plus > 10 or lc.tolerance_minus > 10:
        issues.append(ValidationIssue(
            field="tolerance",
            severity="warning",
            message="Tolerance exceeds 10% which may cause issues with some banks",
            suggestion="Consider using standard 5% or 10% tolerance"
        ))
        risk_factors.append(("high_tolerance", 10))
    
    # Party validations
    if not lc.applicant_name:
        issues.append(ValidationIssue(
            field="applicant_name",
            severity="error",
            message="Applicant name is required"
        ))
    
    if not lc.beneficiary_name:
        issues.append(ValidationIssue(
            field="beneficiary_name",
            severity="error",
            message="Beneficiary name is required"
        ))
    
    # Goods description
    if not lc.goods_description or len(lc.goods_description) < 10:
        issues.append(ValidationIssue(
            field="goods_description",
            severity="warning",
            message="Goods description is too short",
            suggestion="Provide detailed goods description including quantity and specifications"
        ))
    
    # Incoterms + Insurance check
    if lc.incoterms in ["FOB", "CFR", "FCA"]:
        # Check if insurance document is required
        docs = lc.documents_required or []
        has_insurance = any("insurance" in str(d).lower() for d in docs)
        if not has_insurance:
            issues.append(ValidationIssue(
                field="documents_required",
                severity="info",
                message=f"With {lc.incoterms} terms, insurance is buyer's responsibility",
                suggestion="Consider whether insurance certificate should be required from seller"
            ))
    
    # Payment terms validation
    if lc.payment_terms == PaymentTerms.USANCE and not lc.usance_days:
        issues.append(ValidationIssue(
            field="usance_days",
            severity="error",
            message="Usance days required for deferred payment",
            suggestion="Specify number of days (e.g., 30, 60, 90)"
        ))
    
    # Presentation period
    if lc.presentation_period and lc.presentation_period < 14:
        issues.append(ValidationIssue(
            field="presentation_period",
            severity="warning",
            message=f"Presentation period of {lc.presentation_period} days is shorter than typical",
            suggestion="Consider extending to at least 21 days (UCP600 default)"
        ))
        risk_factors.append(("short_presentation", 10))
    
    # Calculate risk score (0-100, lower is better)
    base_score = 0
    for factor, points in risk_factors:
        base_score += points
    
    # Clause-based risk
    if lc.selected_clause_ids:
        for code in lc.selected_clause_ids:
            clause = LCClauseLibrary.get_clause_by_code(code)
            if clause:
                if clause.risk_level == RiskLevel.HIGH:
                    base_score += 15
                elif clause.risk_level == RiskLevel.MEDIUM:
                    base_score += 5
    
    # Error count increases risk
    error_count = len([i for i in issues if i.severity == "error"])
    base_score += error_count * 20
    
    risk_score = min(100, max(0, base_score))
    
    return ValidationResult(
        is_valid=error_count == 0,
        issues=issues,
        risk_score=risk_score,
        risk_details={
            "factors": [f[0] for f in risk_factors],
            "error_count": error_count,
            "warning_count": len([i for i in issues if i.severity == "warning"]),
            "clause_risk_contribution": sum(
                15 if LCClauseLibrary.get_clause_by_code(c) and 
                LCClauseLibrary.get_clause_by_code(c).risk_level == RiskLevel.HIGH else
                5 if LCClauseLibrary.get_clause_by_code(c) and
                LCClauseLibrary.get_clause_by_code(c).risk_level == RiskLevel.MEDIUM else 0
                for c in (lc.selected_clause_ids or [])
            )
        }
    )


# ============================================================================
# LC Application Endpoints
# ============================================================================

@router.post("/applications")
async def create_lc_application(
    data: LCApplicationCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new LC application"""
    
    # Create application
    lc = LCApplication(
        user_id=current_user.id,
        reference_number=generate_reference_number(),
        name=data.name,
        lc_type=LCType(data.lc_type),
        status=LCStatus.DRAFT,
        
        # Amount
        currency=data.currency,
        amount=data.amount,
        tolerance_plus=data.tolerance_plus,
        tolerance_minus=data.tolerance_minus,
        
        # Applicant
        applicant_name=data.applicant.name,
        applicant_address=data.applicant.address,
        applicant_country=data.applicant.country,
        applicant_contact=data.applicant.contact,
        
        # Beneficiary
        beneficiary_name=data.beneficiary.name,
        beneficiary_address=data.beneficiary.address,
        beneficiary_country=data.beneficiary.country,
        beneficiary_contact=data.beneficiary.contact,
        
        # Banks
        issuing_bank_name=data.issuing_bank.name if data.issuing_bank else None,
        issuing_bank_swift=data.issuing_bank.swift if data.issuing_bank else None,
        advising_bank_name=data.advising_bank.name if data.advising_bank else None,
        advising_bank_swift=data.advising_bank.swift if data.advising_bank else None,
        
        # Shipment
        port_of_loading=data.port_of_loading,
        port_of_discharge=data.port_of_discharge,
        place_of_delivery=data.place_of_delivery,
        latest_shipment_date=data.latest_shipment_date,
        incoterms=data.incoterms,
        incoterms_place=data.incoterms_place,
        partial_shipments=data.partial_shipments,
        transhipment=data.transhipment,
        
        # Goods
        goods_description=data.goods_description,
        hs_code=data.hs_code,
        quantity=data.quantity,
        unit_price=data.unit_price,
        
        # Payment
        payment_terms=PaymentTerms(data.payment_terms),
        usance_days=data.usance_days,
        usance_from=data.usance_from,
        
        # Validity
        expiry_date=data.expiry_date,
        expiry_place=data.expiry_place,
        presentation_period=data.presentation_period,
        confirmation_instructions=ConfirmationInstructions(data.confirmation_instructions),
        
        # Documents & Conditions
        documents_required=[d.dict() for d in data.documents_required],
        additional_conditions=data.additional_conditions,
        selected_clause_ids=data.selected_clause_ids,
    )
    
    db.add(lc)
    db.commit()
    db.refresh(lc)
    
    # Validate and update risk score
    validation = validate_lc_application(lc)
    lc.validation_issues = [i.dict() for i in validation.issues]
    lc.risk_score = validation.risk_score
    lc.risk_details = validation.risk_details
    db.commit()
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "status": lc.status.value,
        "validation": validation.dict(),
        "created_at": lc.created_at.isoformat()
    }


@router.get("/applications")
async def list_lc_applications(
    status: Optional[str] = None,
    limit: int = Query(default=20, le=100),
    offset: int = 0,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List user's LC applications"""
    
    query = db.query(LCApplication).filter(
        LCApplication.user_id == current_user.id
    )
    
    if status:
        query = query.filter(LCApplication.status == LCStatus(status))
    
    total = query.count()
    applications = query.order_by(
        LCApplication.updated_at.desc()
    ).offset(offset).limit(limit).all()
    
    return {
        "total": total,
        "applications": [
            {
                "id": str(app.id),
                "reference_number": app.reference_number,
                "name": app.name,
                "status": app.status.value,
                "lc_type": app.lc_type.value,
                "amount": app.amount,
                "currency": app.currency,
                "beneficiary_name": app.beneficiary_name,
                "expiry_date": app.expiry_date.isoformat() if app.expiry_date else None,
                "risk_score": app.risk_score,
                "created_at": app.created_at.isoformat(),
                "updated_at": app.updated_at.isoformat(),
            }
            for app in applications
        ]
    }


@router.get("/applications/{application_id}")
async def get_lc_application(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific LC application"""
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "name": lc.name,
        "status": lc.status.value,
        "lc_type": lc.lc_type.value,
        
        # Amount
        "currency": lc.currency,
        "amount": lc.amount,
        "tolerance_plus": lc.tolerance_plus,
        "tolerance_minus": lc.tolerance_minus,
        
        # Parties
        "applicant": {
            "name": lc.applicant_name,
            "address": lc.applicant_address,
            "country": lc.applicant_country,
            "contact": lc.applicant_contact,
        },
        "beneficiary": {
            "name": lc.beneficiary_name,
            "address": lc.beneficiary_address,
            "country": lc.beneficiary_country,
            "contact": lc.beneficiary_contact,
        },
        "issuing_bank": {
            "name": lc.issuing_bank_name,
            "swift": lc.issuing_bank_swift,
        } if lc.issuing_bank_name else None,
        "advising_bank": {
            "name": lc.advising_bank_name,
            "swift": lc.advising_bank_swift,
        } if lc.advising_bank_name else None,
        
        # Shipment
        "port_of_loading": lc.port_of_loading,
        "port_of_discharge": lc.port_of_discharge,
        "place_of_delivery": lc.place_of_delivery,
        "latest_shipment_date": lc.latest_shipment_date.isoformat() if lc.latest_shipment_date else None,
        "incoterms": lc.incoterms,
        "incoterms_place": lc.incoterms_place,
        "partial_shipments": lc.partial_shipments,
        "transhipment": lc.transhipment,
        
        # Goods
        "goods_description": lc.goods_description,
        "hs_code": lc.hs_code,
        "quantity": lc.quantity,
        "unit_price": lc.unit_price,
        
        # Payment
        "payment_terms": lc.payment_terms.value,
        "usance_days": lc.usance_days,
        "usance_from": lc.usance_from,
        
        # Validity
        "expiry_date": lc.expiry_date.isoformat() if lc.expiry_date else None,
        "expiry_place": lc.expiry_place,
        "presentation_period": lc.presentation_period,
        "confirmation_instructions": lc.confirmation_instructions.value,
        
        # Documents & Conditions
        "documents_required": lc.documents_required,
        "additional_conditions": lc.additional_conditions,
        "selected_clause_ids": lc.selected_clause_ids,
        
        # Validation
        "validation_issues": lc.validation_issues,
        "risk_score": lc.risk_score,
        "risk_details": lc.risk_details,
        
        # Metadata
        "created_at": lc.created_at.isoformat(),
        "updated_at": lc.updated_at.isoformat(),
    }


@router.put("/applications/{application_id}")
async def update_lc_application(
    application_id: str,
    data: LCApplicationUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update an LC application"""
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    if lc.status not in [LCStatus.DRAFT, LCStatus.REVIEW]:
        raise HTTPException(status_code=400, detail="Cannot edit submitted LC application")
    
    # Update fields
    update_data = data.dict(exclude_unset=True)
    for field, value in update_data.items():
        if value is not None:
            if field == "documents_required" and value:
                value = [d.dict() if hasattr(d, 'dict') else d for d in value]
            setattr(lc, field, value)
    
    lc.updated_at = datetime.utcnow()
    
    # Re-validate
    validation = validate_lc_application(lc)
    lc.validation_issues = [i.dict() for i in validation.issues]
    lc.risk_score = validation.risk_score
    lc.risk_details = validation.risk_details
    
    db.commit()
    
    return {
        "id": str(lc.id),
        "status": "updated",
        "validation": validation.dict()
    }


@router.delete("/applications/{application_id}")
async def delete_lc_application(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete an LC application"""
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    if lc.status == LCStatus.SUBMITTED:
        raise HTTPException(status_code=400, detail="Cannot delete submitted LC application")
    
    db.delete(lc)
    db.commit()
    
    return {"status": "deleted"}


@router.post("/applications/{application_id}/validate")
async def validate_application(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Validate an LC application"""
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    validation = validate_lc_application(lc)
    
    # Update stored validation
    lc.validation_issues = [i.dict() for i in validation.issues]
    lc.risk_score = validation.risk_score
    lc.risk_details = validation.risk_details
    db.commit()


# ============================================================================
# Version History Endpoints
# ============================================================================

def _lc_to_snapshot(lc: LCApplication) -> Dict[str, Any]:
    """Convert LC application to a JSON snapshot for versioning"""
    return {
        "reference_number": lc.reference_number,
        "name": lc.name,
        "lc_type": lc.lc_type.value if lc.lc_type else None,
        "status": lc.status.value if lc.status else None,
        "currency": lc.currency,
        "amount": lc.amount,
        "tolerance_plus": lc.tolerance_plus,
        "tolerance_minus": lc.tolerance_minus,
        "applicant_name": lc.applicant_name,
        "applicant_address": lc.applicant_address,
        "applicant_country": lc.applicant_country,
        "beneficiary_name": lc.beneficiary_name,
        "beneficiary_address": lc.beneficiary_address,
        "beneficiary_country": lc.beneficiary_country,
        "issuing_bank_name": lc.issuing_bank_name,
        "issuing_bank_swift": lc.issuing_bank_swift,
        "advising_bank_name": lc.advising_bank_name,
        "advising_bank_swift": lc.advising_bank_swift,
        "port_of_loading": lc.port_of_loading,
        "port_of_discharge": lc.port_of_discharge,
        "place_of_delivery": lc.place_of_delivery,
        "latest_shipment_date": lc.latest_shipment_date.isoformat() if lc.latest_shipment_date else None,
        "incoterms": lc.incoterms,
        "incoterms_place": lc.incoterms_place,
        "partial_shipments": lc.partial_shipments,
        "transhipment": lc.transhipment,
        "goods_description": lc.goods_description,
        "hs_code": lc.hs_code,
        "payment_terms": lc.payment_terms.value if lc.payment_terms else None,
        "usance_days": lc.usance_days,
        "usance_from": lc.usance_from,
        "expiry_date": lc.expiry_date.isoformat() if lc.expiry_date else None,
        "expiry_place": lc.expiry_place,
        "presentation_period": lc.presentation_period,
        "confirmation_instructions": lc.confirmation_instructions.value if lc.confirmation_instructions else None,
        "documents_required": lc.documents_required,
        "additional_conditions": lc.additional_conditions,
        "selected_clause_ids": lc.selected_clause_ids,
    }


def _compute_diff(old_snapshot: Dict[str, Any], new_snapshot: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Compute differences between two snapshots"""
    diffs = []
    
    # Fields to compare with human-readable names
    field_labels = {
        "lc_type": "LC Type",
        "currency": "Currency",
        "amount": "Amount",
        "tolerance_plus": "Tolerance (+%)",
        "tolerance_minus": "Tolerance (-%)",
        "applicant_name": "Applicant Name",
        "applicant_address": "Applicant Address",
        "applicant_country": "Applicant Country",
        "beneficiary_name": "Beneficiary Name",
        "beneficiary_address": "Beneficiary Address",
        "beneficiary_country": "Beneficiary Country",
        "port_of_loading": "Port of Loading",
        "port_of_discharge": "Port of Discharge",
        "place_of_delivery": "Place of Delivery",
        "latest_shipment_date": "Latest Shipment Date",
        "incoterms": "Incoterms",
        "incoterms_place": "Incoterms Place",
        "partial_shipments": "Partial Shipments",
        "transhipment": "Transhipment",
        "goods_description": "Goods Description",
        "hs_code": "HS Code",
        "payment_terms": "Payment Terms",
        "usance_days": "Usance Days",
        "expiry_date": "Expiry Date",
        "expiry_place": "Expiry Place",
        "presentation_period": "Presentation Period",
        "confirmation_instructions": "Confirmation",
    }
    
    for field, label in field_labels.items():
        old_val = old_snapshot.get(field)
        new_val = new_snapshot.get(field)
        
        if old_val != new_val:
            diffs.append({
                "field": field,
                "label": label,
                "old_value": old_val,
                "new_value": new_val,
                "change_type": "modified" if old_val and new_val else ("added" if new_val else "removed")
            })
    
    # Handle complex fields (documents, conditions, clauses)
    old_docs = old_snapshot.get("documents_required", []) or []
    new_docs = new_snapshot.get("documents_required", []) or []
    if old_docs != new_docs:
        diffs.append({
            "field": "documents_required",
            "label": "Documents Required",
            "old_value": f"{len(old_docs)} documents",
            "new_value": f"{len(new_docs)} documents",
            "change_type": "modified"
        })
    
    old_conds = old_snapshot.get("additional_conditions", []) or []
    new_conds = new_snapshot.get("additional_conditions", []) or []
    if old_conds != new_conds:
        diffs.append({
            "field": "additional_conditions",
            "label": "Additional Conditions",
            "old_value": f"{len(old_conds)} conditions",
            "new_value": f"{len(new_conds)} conditions",
            "change_type": "modified"
        })
    
    old_clauses = old_snapshot.get("selected_clause_ids", []) or []
    new_clauses = new_snapshot.get("selected_clause_ids", []) or []
    if old_clauses != new_clauses:
        diffs.append({
            "field": "selected_clause_ids",
            "label": "Selected Clauses",
            "old_value": f"{len(old_clauses)} clauses",
            "new_value": f"{len(new_clauses)} clauses",
            "change_type": "modified"
        })
    
    return diffs


@router.get("/applications/{application_id}/versions")
async def list_versions(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all versions of an LC application"""
    from app.models.lc_builder import LCApplicationVersion
    
    # Verify ownership
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    versions = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.lc_application_id == application_id
    ).order_by(LCApplicationVersion.version_number.desc()).all()
    
    return {
        "application_id": str(application_id),
        "reference_number": lc.reference_number,
        "total_versions": len(versions),
        "versions": [
            {
                "id": str(v.id),
                "version_number": v.version_number,
                "change_summary": v.change_summary,
                "created_at": v.created_at.isoformat(),
                "created_by": str(v.created_by) if v.created_by else None,
            }
            for v in versions
        ]
    }


@router.get("/applications/{application_id}/versions/{version_id}")
async def get_version(
    application_id: str,
    version_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific version snapshot"""
    from app.models.lc_builder import LCApplicationVersion
    
    # Verify ownership
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    version = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.id == version_id,
        LCApplicationVersion.lc_application_id == application_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    
    return {
        "id": str(version.id),
        "application_id": str(application_id),
        "version_number": version.version_number,
        "change_summary": version.change_summary,
        "snapshot": version.snapshot,
        "created_at": version.created_at.isoformat(),
    }


@router.get("/applications/{application_id}/versions/{version_id}/diff")
async def get_version_diff(
    application_id: str,
    version_id: str,
    compare_to: Optional[str] = Query(default=None, description="Version ID to compare against (default: previous version)"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get diff between two versions"""
    from app.models.lc_builder import LCApplicationVersion
    
    # Verify ownership
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Get the target version
    version = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.id == version_id,
        LCApplicationVersion.lc_application_id == application_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    
    # Get comparison version
    if compare_to:
        compare_version = db.query(LCApplicationVersion).filter(
            LCApplicationVersion.id == compare_to,
            LCApplicationVersion.lc_application_id == application_id
        ).first()
    else:
        # Get previous version
        compare_version = db.query(LCApplicationVersion).filter(
            LCApplicationVersion.lc_application_id == application_id,
            LCApplicationVersion.version_number == version.version_number - 1
        ).first()
    
    if compare_version:
        diffs = _compute_diff(compare_version.snapshot, version.snapshot)
        compare_info = {
            "id": str(compare_version.id),
            "version_number": compare_version.version_number,
        }
    else:
        # No previous version - show as all new
        diffs = _compute_diff({}, version.snapshot)
        compare_info = None
    
    return {
        "version": {
            "id": str(version.id),
            "version_number": version.version_number,
            "created_at": version.created_at.isoformat(),
        },
        "compare_to": compare_info,
        "diffs": diffs,
        "change_count": len(diffs),
    }


@router.post("/applications/{application_id}/versions")
async def create_version(
    application_id: str,
    change_summary: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new version snapshot of the current application state"""
    from app.models.lc_builder import LCApplicationVersion
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Get the latest version number
    latest = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.lc_application_id == application_id
    ).order_by(LCApplicationVersion.version_number.desc()).first()
    
    next_version = (latest.version_number + 1) if latest else 1
    
    # Create snapshot
    snapshot = _lc_to_snapshot(lc)
    
    # Auto-generate change summary if not provided
    if not change_summary and latest:
        diffs = _compute_diff(latest.snapshot, snapshot)
        if diffs:
            changed_fields = [d["label"] for d in diffs[:3]]
            change_summary = f"Updated: {', '.join(changed_fields)}"
            if len(diffs) > 3:
                change_summary += f" and {len(diffs) - 3} more"
    
    # Create version record
    version = LCApplicationVersion(
        lc_application_id=lc.id,
        version_number=next_version,
        snapshot=snapshot,
        change_summary=change_summary or f"Version {next_version}",
        created_by=current_user.id
    )
    
    db.add(version)
    db.commit()
    db.refresh(version)
    
    return {
        "id": str(version.id),
        "version_number": version.version_number,
        "change_summary": version.change_summary,
        "created_at": version.created_at.isoformat(),
    }


@router.post("/applications/{application_id}/versions/{version_id}/restore")
async def restore_version(
    application_id: str,
    version_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Restore an LC application to a previous version"""
    from app.models.lc_builder import LCApplicationVersion
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    if lc.status == LCStatus.SUBMITTED:
        raise HTTPException(status_code=400, detail="Cannot restore submitted LC application")
    
    version = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.id == version_id,
        LCApplicationVersion.lc_application_id == application_id
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    
    # Create a backup version of current state first
    current_snapshot = _lc_to_snapshot(lc)
    latest = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.lc_application_id == application_id
    ).order_by(LCApplicationVersion.version_number.desc()).first()
    
    backup_version = LCApplicationVersion(
        lc_application_id=lc.id,
        version_number=(latest.version_number + 1) if latest else 1,
        snapshot=current_snapshot,
        change_summary=f"Backup before restoring to version {version.version_number}",
        created_by=current_user.id
    )
    db.add(backup_version)
    
    # Restore from version snapshot
    snapshot = version.snapshot
    
    lc.name = snapshot.get("name")
    lc.lc_type = LCType(snapshot.get("lc_type")) if snapshot.get("lc_type") else LCType.DOCUMENTARY
    lc.currency = snapshot.get("currency", "USD")
    lc.amount = snapshot.get("amount", 0)
    lc.tolerance_plus = snapshot.get("tolerance_plus", 0)
    lc.tolerance_minus = snapshot.get("tolerance_minus", 0)
    lc.applicant_name = snapshot.get("applicant_name")
    lc.applicant_address = snapshot.get("applicant_address")
    lc.applicant_country = snapshot.get("applicant_country")
    lc.beneficiary_name = snapshot.get("beneficiary_name")
    lc.beneficiary_address = snapshot.get("beneficiary_address")
    lc.beneficiary_country = snapshot.get("beneficiary_country")
    lc.issuing_bank_name = snapshot.get("issuing_bank_name")
    lc.issuing_bank_swift = snapshot.get("issuing_bank_swift")
    lc.advising_bank_name = snapshot.get("advising_bank_name")
    lc.advising_bank_swift = snapshot.get("advising_bank_swift")
    lc.port_of_loading = snapshot.get("port_of_loading")
    lc.port_of_discharge = snapshot.get("port_of_discharge")
    lc.place_of_delivery = snapshot.get("place_of_delivery")
    lc.incoterms = snapshot.get("incoterms")
    lc.incoterms_place = snapshot.get("incoterms_place")
    lc.partial_shipments = snapshot.get("partial_shipments", True)
    lc.transhipment = snapshot.get("transhipment", True)
    lc.goods_description = snapshot.get("goods_description")
    lc.hs_code = snapshot.get("hs_code")
    lc.payment_terms = PaymentTerms(snapshot.get("payment_terms")) if snapshot.get("payment_terms") else PaymentTerms.SIGHT
    lc.usance_days = snapshot.get("usance_days")
    lc.usance_from = snapshot.get("usance_from")
    lc.expiry_place = snapshot.get("expiry_place")
    lc.presentation_period = snapshot.get("presentation_period", 21)
    lc.confirmation_instructions = ConfirmationInstructions(snapshot.get("confirmation_instructions")) if snapshot.get("confirmation_instructions") else ConfirmationInstructions.WITHOUT
    lc.documents_required = snapshot.get("documents_required", [])
    lc.additional_conditions = snapshot.get("additional_conditions", [])
    lc.selected_clause_ids = snapshot.get("selected_clause_ids", [])
    
    # Handle dates
    if snapshot.get("latest_shipment_date"):
        try:
            lc.latest_shipment_date = datetime.fromisoformat(snapshot["latest_shipment_date"].replace("Z", "+00:00"))
        except:
            pass
    
    if snapshot.get("expiry_date"):
        try:
            lc.expiry_date = datetime.fromisoformat(snapshot["expiry_date"].replace("Z", "+00:00"))
        except:
            pass
    
    lc.updated_at = datetime.utcnow()
    
    db.commit()
    
    return {
        "status": "restored",
        "restored_to_version": version.version_number,
        "backup_version": backup_version.version_number,
    }
    
    return validation.dict()


@router.post("/applications/{application_id}/duplicate")
async def duplicate_application(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Duplicate an LC application"""
    
    original = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not original:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Create copy
    new_lc = LCApplication(
        user_id=current_user.id,
        reference_number=generate_reference_number(),
        name=f"Copy of {original.name}" if original.name else None,
        lc_type=original.lc_type,
        status=LCStatus.DRAFT,
        
        currency=original.currency,
        amount=original.amount,
        tolerance_plus=original.tolerance_plus,
        tolerance_minus=original.tolerance_minus,
        
        applicant_name=original.applicant_name,
        applicant_address=original.applicant_address,
        applicant_country=original.applicant_country,
        
        beneficiary_name=original.beneficiary_name,
        beneficiary_address=original.beneficiary_address,
        beneficiary_country=original.beneficiary_country,
        
        port_of_loading=original.port_of_loading,
        port_of_discharge=original.port_of_discharge,
        incoterms=original.incoterms,
        partial_shipments=original.partial_shipments,
        transhipment=original.transhipment,
        
        goods_description=original.goods_description,
        hs_code=original.hs_code,
        
        payment_terms=original.payment_terms,
        usance_days=original.usance_days,
        
        presentation_period=original.presentation_period,
        confirmation_instructions=original.confirmation_instructions,
        
        documents_required=original.documents_required,
        additional_conditions=original.additional_conditions,
        selected_clause_ids=original.selected_clause_ids,
    )
    
    db.add(new_lc)
    db.commit()
    db.refresh(new_lc)
    
    return {
        "id": str(new_lc.id),
        "reference_number": new_lc.reference_number,
        "status": "created"
    }


# ============================================================================
# Amendment Builder
# ============================================================================

class AmendmentRequest(BaseModel):
    """Request to create an amendment"""
    amendment_type: str = Field(..., description="Type: increase_amount, decrease_amount, extend_expiry, change_shipment, add_documents, change_goods, other")
    reason: str = Field(..., description="Reason for the amendment")
    changes: Dict[str, Any] = Field(default_factory=dict, description="Fields to change")


@router.post("/applications/{application_id}/amend")
async def create_amendment(
    application_id: str,
    request: AmendmentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Create an amendment to an existing LC application.
    This creates a new version (if not draft) and applies the changes.
    """
    from app.models.lc_builder import LCApplicationVersion
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Store original state before amendment
    original_snapshot = _lc_to_snapshot(lc)
    
    # Get latest version number
    latest = db.query(LCApplicationVersion).filter(
        LCApplicationVersion.lc_application_id == application_id
    ).order_by(LCApplicationVersion.version_number.desc()).first()
    
    next_version = (latest.version_number + 1) if latest else 1
    
    # Create version record with pre-amendment state
    pre_amendment_version = LCApplicationVersion(
        lc_application_id=lc.id,
        version_number=next_version,
        snapshot=original_snapshot,
        change_summary=f"Pre-amendment snapshot: {request.amendment_type}",
        created_by=current_user.id
    )
    db.add(pre_amendment_version)
    
    # Apply changes based on amendment type
    changes_applied = []
    
    for field, value in request.changes.items():
        if hasattr(lc, field):
            old_value = getattr(lc, field)
            
            # Handle enum fields
            if field == "payment_terms" and isinstance(value, str):
                value = PaymentTerms(value)
            elif field == "confirmation_instructions" and isinstance(value, str):
                value = ConfirmationInstructions(value)
            elif field == "lc_type" and isinstance(value, str):
                value = LCType(value)
            
            # Handle date fields
            if field in ["expiry_date", "latest_shipment_date"] and isinstance(value, str):
                try:
                    from dateutil import parser as date_parser
                    value = date_parser.parse(value)
                except:
                    pass
            
            setattr(lc, field, value)
            changes_applied.append({
                "field": field,
                "old_value": str(old_value) if old_value else None,
                "new_value": str(value) if value else None,
            })
    
    # Update status to AMENDED if it was submitted/approved
    if lc.status in [LCStatus.SUBMITTED, LCStatus.APPROVED]:
        lc.status = LCStatus.AMENDED
    
    lc.updated_at = datetime.utcnow()
    
    # Re-validate
    validation = validate_lc_application(lc)
    lc.validation_issues = [i.dict() for i in validation.issues]
    lc.risk_score = validation.risk_score
    lc.risk_details = validation.risk_details
    
    db.commit()
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "status": lc.status.value,
        "amendment_type": request.amendment_type,
        "changes_applied": changes_applied,
        "version_created": next_version,
        "validation": {
            "is_valid": validation.is_valid,
            "risk_score": validation.risk_score,
            "issue_count": len(validation.issues),
        }
    }


@router.get("/applications/{application_id}/amendment-options")
async def get_amendment_options(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get available amendment options for an LC application.
    Returns common amendment types and which fields can be changed.
    """
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    return {
        "current_values": {
            "amount": lc.amount,
            "currency": lc.currency,
            "expiry_date": lc.expiry_date.isoformat() if lc.expiry_date else None,
            "latest_shipment_date": lc.latest_shipment_date.isoformat() if lc.latest_shipment_date else None,
            "port_of_loading": lc.port_of_loading,
            "port_of_discharge": lc.port_of_discharge,
            "goods_description": lc.goods_description,
            "documents_count": len(lc.documents_required or []),
        },
        "amendment_types": [
            {
                "type": "increase_amount",
                "label": "Increase LC Amount",
                "description": "Increase the documentary credit amount",
                "fields": ["amount"],
                "notes": "May require bank approval",
            },
            {
                "type": "decrease_amount",
                "label": "Decrease LC Amount",
                "description": "Reduce the documentary credit amount",
                "fields": ["amount"],
                "notes": "Usually approved quickly",
            },
            {
                "type": "extend_expiry",
                "label": "Extend Expiry Date",
                "description": "Extend the LC expiry and/or shipment date",
                "fields": ["expiry_date", "latest_shipment_date"],
                "notes": "Common amendment type",
            },
            {
                "type": "change_shipment",
                "label": "Change Shipment Details",
                "description": "Modify ports, partial shipments, or transhipment",
                "fields": ["port_of_loading", "port_of_discharge", "partial_shipments", "transhipment"],
                "notes": "May affect freight costs",
            },
            {
                "type": "add_documents",
                "label": "Add/Remove Documents",
                "description": "Modify required documents list",
                "fields": ["documents_required"],
                "notes": "Check beneficiary capability",
            },
            {
                "type": "change_goods",
                "label": "Change Goods Description",
                "description": "Modify the goods description",
                "fields": ["goods_description", "hs_code"],
                "notes": "Must match actual shipment",
            },
            {
                "type": "change_terms",
                "label": "Change Payment Terms",
                "description": "Modify payment terms or tenor",
                "fields": ["payment_terms", "usance_days"],
                "notes": "May affect pricing",
            },
            {
                "type": "other",
                "label": "Other Amendment",
                "description": "Other changes to LC terms",
                "fields": ["additional_conditions", "tolerance_plus", "tolerance_minus", "incoterms"],
                "notes": "Specify details",
            },
        ]
    }


# ============================================================================
# Profile Management Endpoints (Applicant & Beneficiary)
# ============================================================================

class ApplicantProfileCreate(BaseModel):
    company_name: str
    address: Optional[str] = None
    city: Optional[str] = None
    country: str
    contact_person: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    bank_name: Optional[str] = None
    bank_swift: Optional[str] = None
    bank_address: Optional[str] = None
    is_favorite: bool = False


class BeneficiaryProfileCreate(BaseModel):
    company_name: str
    address: Optional[str] = None
    city: Optional[str] = None
    country: str
    contact_person: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    bank_name: Optional[str] = None
    bank_swift: Optional[str] = None
    bank_account: Optional[str] = None
    bank_address: Optional[str] = None
    industry: Optional[str] = "general"
    is_favorite: bool = False


@router.get("/profiles/applicants")
async def list_applicant_profiles(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List saved applicant profiles"""
    profiles = db.query(ApplicantProfile).filter(
        ApplicantProfile.user_id == current_user.id
    ).order_by(ApplicantProfile.is_favorite.desc(), ApplicantProfile.usage_count.desc()).all()
    
    return {
        "profiles": [
            {
                "id": str(p.id),
                "company_name": p.company_name,
                "address": p.address,
                "city": p.city,
                "country": p.country,
                "contact_person": p.contact_person,
                "email": p.email,
                "phone": p.phone,
                "bank_name": p.bank_name,
                "bank_swift": p.bank_swift,
                "bank_address": p.bank_address,
                "is_favorite": p.is_favorite,
                "usage_count": p.usage_count,
                "created_at": p.created_at.isoformat(),
            }
            for p in profiles
        ]
    }


@router.post("/profiles/applicants")
async def create_applicant_profile(
    data: ApplicantProfileCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new applicant profile"""
    profile = ApplicantProfile(
        user_id=current_user.id,
        company_name=data.company_name,
        address=data.address,
        city=data.city,
        country=data.country,
        contact_person=data.contact_person,
        email=data.email,
        phone=data.phone,
        bank_name=data.bank_name,
        bank_swift=data.bank_swift,
        bank_address=data.bank_address,
        is_favorite=data.is_favorite,
    )
    
    db.add(profile)
    db.commit()
    db.refresh(profile)
    
    return {"id": str(profile.id), "status": "created"}


@router.put("/profiles/applicants/{profile_id}")
async def update_applicant_profile(
    profile_id: str,
    data: ApplicantProfileCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update an applicant profile"""
    profile = db.query(ApplicantProfile).filter(
        ApplicantProfile.id == profile_id,
        ApplicantProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    for field, value in data.dict().items():
        if value is not None:
            setattr(profile, field, value)
    
    db.commit()
    return {"id": str(profile.id), "status": "updated"}


@router.delete("/profiles/applicants/{profile_id}")
async def delete_applicant_profile(
    profile_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete an applicant profile"""
    profile = db.query(ApplicantProfile).filter(
        ApplicantProfile.id == profile_id,
        ApplicantProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    db.delete(profile)
    db.commit()
    return {"status": "deleted"}


@router.get("/profiles/beneficiaries")
async def list_beneficiary_profiles(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List saved beneficiary profiles"""
    profiles = db.query(BeneficiaryProfile).filter(
        BeneficiaryProfile.user_id == current_user.id
    ).order_by(BeneficiaryProfile.is_favorite.desc(), BeneficiaryProfile.usage_count.desc()).all()
    
    return {
        "profiles": [
            {
                "id": str(p.id),
                "company_name": p.company_name,
                "address": p.address,
                "city": p.city,
                "country": p.country,
                "contact_person": p.contact_person,
                "email": p.email,
                "phone": p.phone,
                "bank_name": p.bank_name,
                "bank_swift": p.bank_swift,
                "bank_account": p.bank_account,
                "bank_address": p.bank_address,
                "industry": p.industry,
                "is_favorite": p.is_favorite,
                "usage_count": p.usage_count,
                "created_at": p.created_at.isoformat(),
            }
            for p in profiles
        ]
    }


@router.post("/profiles/beneficiaries")
async def create_beneficiary_profile(
    data: BeneficiaryProfileCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new beneficiary profile"""
    profile = BeneficiaryProfile(
        user_id=current_user.id,
        company_name=data.company_name,
        address=data.address,
        city=data.city,
        country=data.country,
        contact_person=data.contact_person,
        email=data.email,
        phone=data.phone,
        bank_name=data.bank_name,
        bank_swift=data.bank_swift,
        bank_account=data.bank_account,
        bank_address=data.bank_address,
        industry=data.industry,
        is_favorite=data.is_favorite,
    )
    
    db.add(profile)
    db.commit()
    db.refresh(profile)
    
    return {"id": str(profile.id), "status": "created"}


@router.put("/profiles/beneficiaries/{profile_id}")
async def update_beneficiary_profile(
    profile_id: str,
    data: BeneficiaryProfileCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update a beneficiary profile"""
    profile = db.query(BeneficiaryProfile).filter(
        BeneficiaryProfile.id == profile_id,
        BeneficiaryProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    for field, value in data.dict().items():
        if value is not None:
            setattr(profile, field, value)
    
    db.commit()
    return {"id": str(profile.id), "status": "updated"}


@router.delete("/profiles/beneficiaries/{profile_id}")
async def delete_beneficiary_profile(
    profile_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a beneficiary profile"""
    profile = db.query(BeneficiaryProfile).filter(
        BeneficiaryProfile.id == profile_id,
        BeneficiaryProfile.user_id == current_user.id
    ).first()
    
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    db.delete(profile)
    db.commit()
    return {"status": "deleted"}


# ============================================================================
# Import from Previous LC
# ============================================================================

@router.post("/applications/{application_id}/import-data")
async def import_from_previous_lc(
    application_id: str,
    source_id: str,
    fields: List[str] = Query(default=["applicant", "beneficiary", "goods", "documents"]),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Import data from a previous LC application into the current one.
    
    Fields that can be imported:
    - applicant: Applicant name, address, country
    - beneficiary: Beneficiary name, address, country
    - banks: Issuing and advising bank details
    - shipment: Ports, incoterms, partial/transhipment
    - goods: Goods description, HS code
    - documents: Required documents list
    - conditions: Additional conditions
    - clauses: Selected clause IDs
    - payment: Payment terms, tenor
    """
    
    # Get target LC
    target_lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not target_lc:
        raise HTTPException(status_code=404, detail="Target LC application not found")
    
    if target_lc.status != LCStatus.DRAFT:
        raise HTTPException(status_code=400, detail="Can only import into draft LC")
    
    # Get source LC
    source_lc = db.query(LCApplication).filter(
        LCApplication.id == source_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not source_lc:
        raise HTTPException(status_code=404, detail="Source LC application not found")
    
    imported = []
    
    # Import selected fields
    if "applicant" in fields:
        target_lc.applicant_name = source_lc.applicant_name
        target_lc.applicant_address = source_lc.applicant_address
        target_lc.applicant_country = source_lc.applicant_country
        target_lc.applicant_contact = source_lc.applicant_contact
        imported.append("applicant")
    
    if "beneficiary" in fields:
        target_lc.beneficiary_name = source_lc.beneficiary_name
        target_lc.beneficiary_address = source_lc.beneficiary_address
        target_lc.beneficiary_country = source_lc.beneficiary_country
        target_lc.beneficiary_contact = source_lc.beneficiary_contact
        imported.append("beneficiary")
    
    if "banks" in fields:
        target_lc.issuing_bank_name = source_lc.issuing_bank_name
        target_lc.issuing_bank_swift = source_lc.issuing_bank_swift
        target_lc.advising_bank_name = source_lc.advising_bank_name
        target_lc.advising_bank_swift = source_lc.advising_bank_swift
        imported.append("banks")
    
    if "shipment" in fields:
        target_lc.port_of_loading = source_lc.port_of_loading
        target_lc.port_of_discharge = source_lc.port_of_discharge
        target_lc.place_of_delivery = source_lc.place_of_delivery
        target_lc.incoterms = source_lc.incoterms
        target_lc.incoterms_place = source_lc.incoterms_place
        target_lc.partial_shipments = source_lc.partial_shipments
        target_lc.transhipment = source_lc.transhipment
        imported.append("shipment")
    
    if "goods" in fields:
        target_lc.goods_description = source_lc.goods_description
        target_lc.hs_code = source_lc.hs_code
        imported.append("goods")
    
    if "documents" in fields:
        target_lc.documents_required = source_lc.documents_required
        imported.append("documents")
    
    if "conditions" in fields:
        target_lc.additional_conditions = source_lc.additional_conditions
        imported.append("conditions")
    
    if "clauses" in fields:
        target_lc.selected_clause_ids = source_lc.selected_clause_ids
        imported.append("clauses")
    
    if "payment" in fields:
        target_lc.payment_terms = source_lc.payment_terms
        target_lc.usance_days = source_lc.usance_days
        target_lc.usance_from = source_lc.usance_from
        target_lc.presentation_period = source_lc.presentation_period
        target_lc.confirmation_instructions = source_lc.confirmation_instructions
        imported.append("payment")
    
    target_lc.updated_at = datetime.utcnow()
    db.commit()
    
    return {
        "status": "imported",
        "fields_imported": imported,
        "source_reference": source_lc.reference_number,
        "target_id": str(target_lc.id)
    }


# ============================================================================
# LCopilot Integration - Import from Validation Sessions
# ============================================================================

@router.get("/lcopilot/sessions")
async def list_lcopilot_sessions(
    limit: int = Query(default=20, le=50),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    List user's LCopilot validation sessions that can be imported.
    Returns sessions with extracted LC data.
    """
    from app.models import ValidationSession, SessionStatus
    
    # Get user's completed validation sessions
    sessions = db.query(ValidationSession).filter(
        ValidationSession.user_id == current_user.id,
        ValidationSession.status.in_([SessionStatus.COMPLETED.value, "completed", "validated"]),
        ValidationSession.extracted_data.isnot(None)
    ).order_by(ValidationSession.created_at.desc()).limit(limit).all()
    
    result = []
    for s in sessions:
        extracted = s.extracted_data or {}
        lc_data = extracted.get("lc") or {}
        mt700 = lc_data.get("mt700") or {}
        
        result.append({
            "session_id": str(s.id),
            "created_at": s.created_at.isoformat() if s.created_at else None,
            "status": s.status,
            "lc_number": lc_data.get("lc_number") or mt700.get("20") or "Unknown",
            "beneficiary": lc_data.get("beneficiary_name") or mt700.get("beneficiary") or "",
            "applicant": lc_data.get("applicant_name") or mt700.get("applicant") or "",
            "amount": lc_data.get("amount") or mt700.get("32B_amount"),
            "currency": lc_data.get("currency") or mt700.get("32B_currency") or "USD",
            "has_extracted_data": bool(lc_data),
        })
    
    return {"sessions": result, "count": len(result)}


@router.get("/lcopilot/sessions/{session_id}/preview")
async def preview_lcopilot_session(
    session_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Preview the LC data extracted from an LCopilot validation session.
    Shows what fields will be imported.
    """
    from app.models import ValidationSession
    
    session = db.query(ValidationSession).filter(
        ValidationSession.id == session_id,
        ValidationSession.user_id == current_user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="Validation session not found")
    
    if not session.extracted_data:
        raise HTTPException(status_code=400, detail="Session has no extracted data")
    
    extracted = session.extracted_data or {}
    lc_data = extracted.get("lc") or {}
    mt700 = lc_data.get("mt700") or {}
    
    # Build preview of importable fields
    preview = {
        "session_id": str(session.id),
        "created_at": session.created_at.isoformat() if session.created_at else None,
        
        # Basic LC Info
        "lc_number": lc_data.get("lc_number") or mt700.get("20"),
        "lc_type": lc_data.get("lc_type") or lc_data.get("form_of_doc_credit") or mt700.get("form_of_doc_credit"),
        
        # Amount
        "currency": lc_data.get("currency") or mt700.get("32B_currency"),
        "amount": lc_data.get("amount") or mt700.get("32B_amount"),
        "tolerance_plus": lc_data.get("tolerance_plus") or mt700.get("39A_tolerance_plus"),
        "tolerance_minus": lc_data.get("tolerance_minus") or mt700.get("39A_tolerance_minus"),
        
        # Parties
        "applicant_name": lc_data.get("applicant_name") or lc_data.get("applicant") or mt700.get("50"),
        "applicant_address": lc_data.get("applicant_address") or mt700.get("applicant_address"),
        "beneficiary_name": lc_data.get("beneficiary_name") or lc_data.get("beneficiary") or mt700.get("59"),
        "beneficiary_address": lc_data.get("beneficiary_address") or mt700.get("beneficiary_address"),
        
        # Banks
        "issuing_bank": lc_data.get("issuing_bank") or mt700.get("issuing_bank"),
        "advising_bank": lc_data.get("advising_bank") or mt700.get("advising_bank"),
        
        # Shipment
        "port_of_loading": lc_data.get("port_of_loading") or mt700.get("44E"),
        "port_of_discharge": lc_data.get("port_of_discharge") or mt700.get("44F"),
        "latest_shipment_date": lc_data.get("latest_shipment_date") or mt700.get("44C"),
        "incoterms": lc_data.get("incoterms"),
        "partial_shipments": lc_data.get("partial_shipments") or mt700.get("43P"),
        "transhipment": lc_data.get("transhipment") or mt700.get("43T"),
        
        # Goods
        "goods_description": lc_data.get("goods_description") or mt700.get("45A"),
        
        # Validity
        "expiry_date": lc_data.get("expiry_date") or mt700.get("31D_expiry_date"),
        "expiry_place": lc_data.get("expiry_place") or mt700.get("31D_expiry_place"),
        "presentation_period": lc_data.get("presentation_period") or mt700.get("48"),
        
        # Conditions
        "additional_conditions": lc_data.get("additional_conditions") or lc_data.get("clauses") or mt700.get("47A") or [],
        
        # Documents
        "documents_required": lc_data.get("documents_required") or mt700.get("46A") or [],
    }
    
    # Count non-null fields
    fields_available = sum(1 for v in preview.values() if v is not None and v != "" and v != [])
    
    return {
        "preview": preview,
        "fields_available": fields_available,
        "can_import": fields_available >= 3,  # At least 3 meaningful fields
    }


@router.post("/lcopilot/import/{session_id}")
async def import_from_lcopilot(
    session_id: str,
    name: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Create a new LC application from an LCopilot validation session.
    Imports all extracted LC data.
    """
    from app.models import ValidationSession
    
    session = db.query(ValidationSession).filter(
        ValidationSession.id == session_id,
        ValidationSession.user_id == current_user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="Validation session not found")
    
    if not session.extracted_data:
        raise HTTPException(status_code=400, detail="Session has no extracted data")
    
    extracted = session.extracted_data or {}
    lc_data = extracted.get("lc") or {}
    mt700 = lc_data.get("mt700") or {}
    
    # Parse amount
    amount_raw = lc_data.get("amount") or mt700.get("32B_amount") or 0
    try:
        if isinstance(amount_raw, str):
            amount = float(amount_raw.replace(",", ""))
        elif isinstance(amount_raw, dict):
            amount = float(amount_raw.get("value", 0))
        else:
            amount = float(amount_raw)
    except (ValueError, TypeError):
        amount = 0.0
    
    # Parse expiry date
    expiry_str = lc_data.get("expiry_date") or mt700.get("31D_expiry_date")
    expiry_date = None
    if expiry_str:
        try:
            from dateutil import parser as date_parser
            expiry_date = date_parser.parse(expiry_str)
        except:
            # Default to 90 days from now
            expiry_date = datetime.utcnow() + timedelta(days=90)
    else:
        expiry_date = datetime.utcnow() + timedelta(days=90)
    
    # Determine LC type
    lc_type_str = (lc_data.get("lc_type") or lc_data.get("form_of_doc_credit") or 
                   mt700.get("form_of_doc_credit") or "documentary").lower()
    
    if "standby" in lc_type_str or "sblc" in lc_type_str:
        lc_type = LCType.STANDBY
    elif "revolving" in lc_type_str:
        lc_type = LCType.REVOLVING
    elif "transferable" in lc_type_str:
        lc_type = LCType.TRANSFERABLE
    else:
        lc_type = LCType.DOCUMENTARY
    
    # Parse tolerances
    tolerance_plus = lc_data.get("tolerance_plus") or mt700.get("39A_tolerance_plus") or 5
    tolerance_minus = lc_data.get("tolerance_minus") or mt700.get("39A_tolerance_minus") or 5
    try:
        tolerance_plus = float(str(tolerance_plus).replace("%", ""))
        tolerance_minus = float(str(tolerance_minus).replace("%", ""))
    except:
        tolerance_plus = 5.0
        tolerance_minus = 5.0
    
    # Parse partial shipments / transhipment
    partial_raw = lc_data.get("partial_shipments") or mt700.get("43P")
    tranship_raw = lc_data.get("transhipment") or mt700.get("43T")
    
    partial_shipments = True
    if partial_raw:
        partial_str = str(partial_raw).lower()
        partial_shipments = "allowed" in partial_str or "permitted" in partial_str or partial_str == "true"
    
    transhipment = True
    if tranship_raw:
        tranship_str = str(tranship_raw).lower()
        transhipment = "allowed" in tranship_str or "permitted" in tranship_str or tranship_str == "true"
    
    # Build documents list
    docs_raw = lc_data.get("documents_required") or mt700.get("46A") or []
    documents_required = []
    if isinstance(docs_raw, list):
        for doc in docs_raw:
            if isinstance(doc, str):
                documents_required.append({
                    "document_type": doc,
                    "copies_original": 1,
                    "copies_copy": 0,
                    "is_required": True,
                })
            elif isinstance(doc, dict):
                documents_required.append(doc)
    
    # Create LC Application
    lc = LCApplication(
        user_id=current_user.id,
        reference_number=generate_reference_number(),
        name=name or f"Import from LCopilot - {session.created_at.strftime('%Y-%m-%d') if session.created_at else 'Session'}",
        lc_type=lc_type,
        status=LCStatus.DRAFT,
        
        # Amount
        currency=lc_data.get("currency") or mt700.get("32B_currency") or "USD",
        amount=amount,
        tolerance_plus=tolerance_plus,
        tolerance_minus=tolerance_minus,
        
        # Parties
        applicant_name=lc_data.get("applicant_name") or lc_data.get("applicant") or mt700.get("50") or "Unknown Applicant",
        applicant_address=lc_data.get("applicant_address") or mt700.get("applicant_address"),
        applicant_country=lc_data.get("applicant_country"),
        
        beneficiary_name=lc_data.get("beneficiary_name") or lc_data.get("beneficiary") or mt700.get("59") or "Unknown Beneficiary",
        beneficiary_address=lc_data.get("beneficiary_address") or mt700.get("beneficiary_address"),
        beneficiary_country=lc_data.get("beneficiary_country"),
        
        # Banks
        issuing_bank_name=lc_data.get("issuing_bank") or mt700.get("issuing_bank"),
        advising_bank_name=lc_data.get("advising_bank") or mt700.get("advising_bank"),
        
        # Shipment
        port_of_loading=lc_data.get("port_of_loading") or mt700.get("44E"),
        port_of_discharge=lc_data.get("port_of_discharge") or mt700.get("44F"),
        place_of_delivery=lc_data.get("place_of_delivery"),
        incoterms=lc_data.get("incoterms"),
        partial_shipments=partial_shipments,
        transhipment=transhipment,
        
        # Goods
        goods_description=lc_data.get("goods_description") or mt700.get("45A") or "Imported from LCopilot",
        
        # Validity
        expiry_date=expiry_date,
        expiry_place=lc_data.get("expiry_place") or mt700.get("31D_expiry_place"),
        presentation_period=lc_data.get("presentation_period") or mt700.get("48") or 21,
        
        # Documents
        documents_required=documents_required,
        
        # Additional conditions
        additional_conditions=lc_data.get("additional_conditions") or lc_data.get("clauses") or [],
    )
    
    db.add(lc)
    db.commit()
    db.refresh(lc)
    
    # Validate and get risk score
    validation = validate_lc_application(lc)
    lc.validation_issues = [i.dict() for i in validation.issues]
    lc.risk_score = validation.risk_score
    lc.risk_details = validation.risk_details
    db.commit()
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "status": lc.status.value,
        "imported_from": str(session.id),
        "validation": {
            "is_valid": validation.is_valid,
            "risk_score": validation.risk_score,
            "issue_count": len(validation.issues),
        },
        "created_at": lc.created_at.isoformat()
    }


# ============================================================================
# Clause Library Endpoints
# ============================================================================

@router.get("/clauses")
async def list_clauses(
    category: Optional[str] = None,
    risk_level: Optional[str] = None,
    bias: Optional[str] = None,
    search: Optional[str] = None,
    limit: int = Query(default=500, le=1000),
    offset: int = 0
):
    """List clauses from the library"""
    
    cat = ClauseCategory(category) if category else None
    risk = RiskLevel(risk_level) if risk_level else None
    bi = BiasIndicator(bias) if bias else None
    
    clauses = LCClauseLibrary.search_clauses(
        query=search or "",
        category=cat,
        risk_level=risk,
        bias=bi
    )
    
    # Get statistics
    stats = LCClauseLibrary.get_statistics()
    
    total = len(clauses)
    paged = clauses[offset:offset + limit]
    
    return {
        "total": total,
        "clauses": [
            {
                "code": c.code,
                "category": c.category.value,
                "subcategory": c.subcategory,
                "title": c.title,
                "clause_text": c.clause_text,
                "plain_english": c.plain_english,
                "risk_level": c.risk_level.value,
                "bias": c.bias.value,
                "risk_notes": c.risk_notes,
                "bank_acceptance": c.bank_acceptance,
                "tags": c.tags,
            }
            for c in paged
        ],
        "statistics": {
            "total_clauses": stats["total_clauses"],
            "categories": stats["categories"],
            "by_risk_level": stats["by_risk_level"],
            "by_bias": stats["by_bias"],
        }
    }


@router.get("/clauses/categories")
async def get_clause_categories():
    """Get clause category counts"""
    stats = LCClauseLibrary.get_statistics()
    return {
        "total": stats["total_clauses"],
        "categories": stats["categories"],
        "by_risk_level": stats["by_risk_level"],
        "by_bias": stats["by_bias"],
    }


@router.get("/clauses/{code}")
async def get_clause(code: str):
    """Get a specific clause by code"""
    
    clause = LCClauseLibrary.get_clause_by_code(code)
    if not clause:
        raise HTTPException(status_code=404, detail="Clause not found")
    
    return {
        "code": clause.code,
        "category": clause.category.value,
        "subcategory": clause.subcategory,
        "title": clause.title,
        "clause_text": clause.clause_text,
        "plain_english": clause.plain_english,
        "risk_level": clause.risk_level.value,
        "bias": clause.bias.value,
        "risk_notes": clause.risk_notes,
        "bank_acceptance": clause.bank_acceptance,
        "tags": clause.tags,
    }


@router.get("/clauses/suggest")
async def suggest_clauses(
    origin_country: Optional[str] = None,
    destination_country: Optional[str] = None,
    goods_type: Optional[str] = None,
    payment_terms: Optional[str] = None,
    incoterms: Optional[str] = None,
    first_time: bool = False,
    amount: Optional[float] = None
):
    """
    Get smart clause suggestions based on trade parameters.
    
    Parameters:
    - origin_country: Country where goods originate (e.g., "Bangladesh", "China")
    - destination_country: Country where goods are shipped to (e.g., "USA", "Germany")
    - goods_type: Type of goods (e.g., "textiles", "electronics", "machinery")
    - payment_terms: Payment terms (e.g., "sight", "usance")
    - incoterms: Incoterms (e.g., "FOB", "CIF", "CFR")
    - first_time: Whether this is a first-time transaction with the beneficiary
    - amount: LC amount in USD (for risk-based suggestions)
    """
    
    clauses = LCClauseLibrary.suggest_clauses(
        origin_country=origin_country,
        destination_country=destination_country,
        goods_type=goods_type,
        payment_terms=payment_terms,
        incoterms=incoterms,
        first_time_beneficiary=first_time,
        amount_usd=amount
    )
    
    return {
        "suggestions": [
            {
                "code": c.code,
                "category": c.category.value,
                "subcategory": c.subcategory,
                "title": c.title,
                "clause_text": c.clause_text,
                "plain_english": c.plain_english,
                "risk_level": c.risk_level.value,
                "bias": c.bias.value,
                "tags": c.tags,
            }
            for c in clauses
        ],
        "count": len(clauses),
        "parameters": {
            "origin_country": origin_country,
            "destination_country": destination_country,
            "goods_type": goods_type,
            "payment_terms": payment_terms,
            "incoterms": incoterms,
            "first_time": first_time,
            "amount": amount,
        }
    }


# ============================================================================
# Template Endpoints
# ============================================================================

@router.get("/templates")
async def list_templates(
    trade_route: Optional[str] = None,
    industry: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """List available templates"""
    
    query = db.query(LCTemplate)
    
    # Include public/system templates
    query = query.filter(
        (LCTemplate.is_public == True) |
        (LCTemplate.is_system == True)
    )
    
    if trade_route:
        query = query.filter(LCTemplate.trade_route.ilike(f"%{trade_route}%"))
    
    if industry:
        query = query.filter(LCTemplate.industry.ilike(f"%{industry}%"))
    
    templates = query.all()
    
    return {
        "templates": [
            {
                "id": str(t.id),
                "name": t.name,
                "description": t.description,
                "trade_route": t.trade_route,
                "industry": t.industry,
                "is_system": t.is_system,
                "usage_count": t.usage_count,
            }
            for t in templates
        ]
    }


@router.get("/templates/{template_id}")
async def get_template(
    template_id: str,
    db: Session = Depends(get_db)
):
    """Get template details"""
    
    template = db.query(LCTemplate).filter(LCTemplate.id == template_id).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    return {
        "id": str(template.id),
        "name": template.name,
        "description": template.description,
        "trade_route": template.trade_route,
        "industry": template.industry,
        "template_data": template.template_data,
        "default_clause_ids": template.default_clause_ids,
        "default_documents": template.default_documents,
    }


class TemplateCreate(BaseModel):
    name: str
    description: Optional[str] = None
    trade_route: Optional[str] = None
    industry: Optional[str] = None
    template_data: Dict[str, Any] = Field(default_factory=dict)
    default_clause_ids: List[str] = Field(default_factory=list)
    default_documents: List[Dict[str, Any]] = Field(default_factory=list)
    is_public: bool = False


@router.post("/templates")
async def create_template(
    data: TemplateCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new LC template"""
    
    template = LCTemplate(
        user_id=current_user.id,
        name=data.name,
        description=data.description,
        trade_route=data.trade_route,
        industry=data.industry,
        template_data=data.template_data,
        default_clause_ids=data.default_clause_ids,
        default_documents=data.default_documents,
        is_public=data.is_public,
        is_system=False,
    )
    
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"id": str(template.id), "status": "created"}


@router.put("/templates/{template_id}")
async def update_template(
    template_id: str,
    data: TemplateCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update a user's template"""
    
    template = db.query(LCTemplate).filter(
        LCTemplate.id == template_id,
        LCTemplate.user_id == current_user.id
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found or not owned by user")
    
    template.name = data.name
    template.description = data.description
    template.trade_route = data.trade_route
    template.industry = data.industry
    template.template_data = data.template_data
    template.default_clause_ids = data.default_clause_ids
    template.default_documents = data.default_documents
    template.is_public = data.is_public
    
    db.commit()
    
    return {"id": str(template.id), "status": "updated"}


@router.delete("/templates/{template_id}")
async def delete_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a user's template"""
    
    template = db.query(LCTemplate).filter(
        LCTemplate.id == template_id,
        LCTemplate.user_id == current_user.id,
        LCTemplate.is_system == False  # Cannot delete system templates
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="Template not found or cannot be deleted")
    
    db.delete(template)
    db.commit()
    
    return {"status": "deleted"}


@router.post("/applications/{application_id}/save-as-template")
async def save_lc_as_template(
    application_id: str,
    name: str,
    description: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Save an LC application as a reusable template"""
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Build trade route from countries
    trade_route = None
    if lc.beneficiary_country and lc.applicant_country:
        trade_route = f"{lc.beneficiary_country} → {lc.applicant_country}"
    
    # Create template data from the LC
    template_data = {
        "lc_type": lc.lc_type.value if lc.lc_type else None,
        "currency": lc.currency,
        "tolerance_plus": lc.tolerance_plus,
        "tolerance_minus": lc.tolerance_minus,
        "incoterms": lc.incoterms,
        "incoterms_place": lc.incoterms_place,
        "partial_shipments": lc.partial_shipments,
        "transhipment": lc.transhipment,
        "payment_terms": lc.payment_terms.value if lc.payment_terms else None,
        "usance_days": lc.usance_days,
        "usance_from": lc.usance_from,
        "presentation_period": lc.presentation_period,
        "confirmation_instructions": lc.confirmation_instructions.value if lc.confirmation_instructions else None,
    }
    
    template = LCTemplate(
        user_id=current_user.id,
        name=name,
        description=description or f"Created from LC {lc.reference_number}",
        trade_route=trade_route,
        industry=None,
        template_data=template_data,
        default_clause_ids=lc.selected_clause_ids or [],
        default_documents=lc.documents_required or [],
        additional_conditions=lc.additional_conditions or [],
        is_public=False,
        is_system=False,
    )
    
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {
        "id": str(template.id),
        "name": template.name,
        "trade_route": trade_route,
        "status": "created"
    }


@router.get("/templates/mine")
async def list_user_templates(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List templates created by the current user"""
    
    templates = db.query(LCTemplate).filter(
        LCTemplate.user_id == current_user.id
    ).order_by(LCTemplate.usage_count.desc()).all()
    
    return {
        "templates": [
            {
                "id": str(t.id),
                "name": t.name,
                "description": t.description,
                "trade_route": t.trade_route,
                "industry": t.industry,
                "usage_count": t.usage_count,
                "is_public": t.is_public,
                "created_at": t.created_at.isoformat() if t.created_at else None,
            }
            for t in templates
        ]
    }


# ============================================================================
# Profile Endpoints
# ============================================================================

@router.get("/profiles/applicants")
async def list_applicant_profiles(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List saved applicant profiles"""
    
    profiles = db.query(ApplicantProfile).filter(
        ApplicantProfile.user_id == current_user.id
    ).all()
    
    return {
        "profiles": [
            {
                "id": str(p.id),
                "name": p.name,
                "address": p.address,
                "country": p.country,
                "is_default": p.is_default,
            }
            for p in profiles
        ]
    }


@router.get("/profiles/beneficiaries")
async def list_beneficiary_profiles(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List saved beneficiary profiles"""
    
    profiles = db.query(BeneficiaryProfile).filter(
        BeneficiaryProfile.user_id == current_user.id
    ).all()
    
    return {
        "profiles": [
            {
                "id": str(p.id),
                "name": p.name,
                "address": p.address,
                "country": p.country,
                "trade_history_count": p.trade_history_count,
            }
            for p in profiles
        ]
    }


# ============================================================================
# Export Endpoints
# ============================================================================

@router.post("/applications/{application_id}/export/mt700")
async def export_mt700(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export LC application as MT700 SWIFT message"""
    from app.services.mt700_generator import MT700Generator
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Build LC data dict
    lc_data = {
        "reference_number": lc.reference_number,
        "lc_type": lc.lc_type.value,
        "currency": lc.currency,
        "amount": lc.amount,
        "tolerance_plus": lc.tolerance_plus,
        "tolerance_minus": lc.tolerance_minus,
        "applicant": {
            "name": lc.applicant_name,
            "address": lc.applicant_address,
            "country": lc.applicant_country,
        },
        "beneficiary": {
            "name": lc.beneficiary_name,
            "address": lc.beneficiary_address,
            "country": lc.beneficiary_country,
        },
        "port_of_loading": lc.port_of_loading,
        "port_of_discharge": lc.port_of_discharge,
        "place_of_delivery": lc.place_of_delivery,
        "latest_shipment_date": lc.latest_shipment_date,
        "incoterms": lc.incoterms,
        "incoterms_place": lc.incoterms_place,
        "partial_shipments": lc.partial_shipments,
        "transhipment": lc.transhipment,
        "goods_description": lc.goods_description,
        "payment_terms": lc.payment_terms.value,
        "usance_days": lc.usance_days,
        "usance_from": lc.usance_from,
        "expiry_date": lc.expiry_date,
        "expiry_place": lc.expiry_place,
        "presentation_period": lc.presentation_period,
        "confirmation_instructions": lc.confirmation_instructions.value,
        "documents_required": lc.documents_required,
        "additional_conditions": lc.additional_conditions,
    }
    
    generator = MT700Generator(lc_data)
    result = generator.build()
    
    return result


@router.get("/mt700/fields")
async def get_mt700_field_reference():
    """Get MT700 field reference for UI"""
    from app.services.mt700_generator import MT700Generator
    return {"fields": MT700Generator.field_reference()}


@router.post("/applications/{application_id}/export/pdf")
async def export_pdf(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export LC application as PDF"""
    from fastapi.responses import Response
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Generate PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
        alignment=1  # Center
    )
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=15,
        spaceAfter=10,
        textColor=colors.HexColor("#1e40af")
    )
    normal_style = styles['Normal']
    
    elements = []
    
    # Title
    elements.append(Paragraph("DOCUMENTARY CREDIT APPLICATION", title_style))
    elements.append(Paragraph(f"Reference: {lc.reference_number}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Basic Info
    elements.append(Paragraph("1. BASIC INFORMATION", heading_style))
    basic_data = [
        ["Type of Credit:", lc.lc_type.value.upper()],
        ["Currency & Amount:", f"{lc.currency} {lc.amount:,.2f}"],
        ["Tolerance:", f"+{lc.tolerance_plus}% / -{lc.tolerance_minus}%"],
    ]
    basic_table = Table(basic_data, colWidths=[2*inch, 4*inch])
    basic_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(basic_table)
    
    # Parties
    elements.append(Paragraph("2. PARTIES", heading_style))
    parties_data = [
        ["APPLICANT:", "BENEFICIARY:"],
        [lc.applicant_name or "", lc.beneficiary_name or ""],
        [lc.applicant_address or "", lc.beneficiary_address or ""],
        [lc.applicant_country or "", lc.beneficiary_country or ""],
    ]
    parties_table = Table(parties_data, colWidths=[3*inch, 3*inch])
    parties_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(parties_table)
    
    # Shipment
    elements.append(Paragraph("3. SHIPMENT DETAILS", heading_style))
    ship_data = [
        ["Port of Loading:", lc.port_of_loading or "Any Port"],
        ["Port of Discharge:", lc.port_of_discharge or ""],
        ["Latest Shipment:", lc.latest_shipment_date.strftime("%d %B %Y") if lc.latest_shipment_date else ""],
        ["Incoterms:", f"{lc.incoterms} {lc.incoterms_place or ''}" if lc.incoterms else ""],
        ["Partial Shipments:", "ALLOWED" if lc.partial_shipments else "NOT ALLOWED"],
        ["Transhipment:", "ALLOWED" if lc.transhipment else "NOT ALLOWED"],
    ]
    ship_table = Table(ship_data, colWidths=[2*inch, 4*inch])
    ship_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    elements.append(ship_table)
    
    # Goods
    elements.append(Paragraph("4. GOODS DESCRIPTION", heading_style))
    elements.append(Paragraph(lc.goods_description or "", normal_style))
    
    # Payment
    elements.append(Paragraph("5. PAYMENT TERMS", heading_style))
    payment_text = lc.payment_terms.value.upper()
    if lc.payment_terms.value == "usance" and lc.usance_days:
        payment_text = f"{lc.usance_days} DAYS FROM {lc.usance_from or 'B/L DATE'}"
    pay_data = [
        ["Payment:", payment_text],
        ["Expiry Date:", lc.expiry_date.strftime("%d %B %Y") if lc.expiry_date else ""],
        ["Expiry Place:", lc.expiry_place or ""],
        ["Presentation Period:", f"{lc.presentation_period} days after shipment"],
        ["Confirmation:", lc.confirmation_instructions.value.upper()],
    ]
    pay_table = Table(pay_data, colWidths=[2*inch, 4*inch])
    pay_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    elements.append(pay_table)
    
    # Documents
    elements.append(Paragraph("6. DOCUMENTS REQUIRED", heading_style))
    if lc.documents_required:
        for i, doc in enumerate(lc.documents_required, 1):
            if isinstance(doc, dict):
                doc_text = f"{i}. {doc.get('document_type', '')} - {doc.get('copies_original', 1)} original(s)"
                if doc.get('specific_requirements'):
                    doc_text += f"\n   {doc['specific_requirements']}"
            else:
                doc_text = f"{i}. {doc}"
            elements.append(Paragraph(doc_text, normal_style))
    
    # Additional Conditions
    if lc.additional_conditions:
        elements.append(Paragraph("7. ADDITIONAL CONDITIONS", heading_style))
        for cond in lc.additional_conditions:
            elements.append(Paragraph(f"• {cond}", normal_style))
    
    # Footer
    elements.append(Spacer(1, 30))
    elements.append(Paragraph(
        f"Generated on {datetime.now().strftime('%d %B %Y at %H:%M')} by TRDR Hub LC Builder",
        ParagraphStyle('Footer', parent=normal_style, fontSize=8, textColor=colors.gray)
    ))
    
    doc.build(elements)
    
    buffer.seek(0)
    pdf_bytes = buffer.getvalue()
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=LC_{lc.reference_number}.pdf"
        }
    )


@router.post("/applications/{application_id}/export/word")
async def export_word(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export LC application as Word document (.docx)"""
    from fastapi.responses import Response
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Create Word document
    doc = Document()
    
    # Set document style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Title
    title = doc.add_heading('DOCUMENTARY CREDIT APPLICATION', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Reference
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Reference: {lc.reference_number}")
    ref_run.bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Helper function to add section heading
    def add_section_heading(text):
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(30, 64, 175)  # Blue color
    
    # 1. Basic Information
    add_section_heading("1. BASIC INFORMATION")
    
    basic_table = doc.add_table(rows=3, cols=2)
    basic_table.style = 'Table Grid'
    
    basic_data = [
        ("Type of Credit:", lc.lc_type.value.upper()),
        ("Currency & Amount:", f"{lc.currency} {lc.amount:,.2f}"),
        ("Tolerance:", f"+{lc.tolerance_plus}% / -{lc.tolerance_minus}%"),
    ]
    
    for i, (label, value) in enumerate(basic_data):
        row = basic_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # 2. Parties
    add_section_heading("2. PARTIES")
    
    parties_table = doc.add_table(rows=4, cols=2)
    parties_table.style = 'Table Grid'
    
    # Headers
    parties_table.rows[0].cells[0].text = "APPLICANT"
    parties_table.rows[0].cells[1].text = "BENEFICIARY"
    parties_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    parties_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    # Data
    parties_table.rows[1].cells[0].text = lc.applicant_name or ""
    parties_table.rows[1].cells[1].text = lc.beneficiary_name or ""
    parties_table.rows[2].cells[0].text = lc.applicant_address or ""
    parties_table.rows[2].cells[1].text = lc.beneficiary_address or ""
    parties_table.rows[3].cells[0].text = lc.applicant_country or ""
    parties_table.rows[3].cells[1].text = lc.beneficiary_country or ""
    
    doc.add_paragraph()  # Spacer
    
    # 3. Shipment Details
    add_section_heading("3. SHIPMENT DETAILS")
    
    ship_table = doc.add_table(rows=6, cols=2)
    ship_table.style = 'Table Grid'
    
    ship_data = [
        ("Port of Loading:", lc.port_of_loading or "Any Port"),
        ("Port of Discharge:", lc.port_of_discharge or ""),
        ("Latest Shipment:", lc.latest_shipment_date.strftime("%d %B %Y") if lc.latest_shipment_date else ""),
        ("Incoterms:", f"{lc.incoterms} {lc.incoterms_place or ''}" if lc.incoterms else ""),
        ("Partial Shipments:", "ALLOWED" if lc.partial_shipments else "NOT ALLOWED"),
        ("Transhipment:", "ALLOWED" if lc.transhipment else "NOT ALLOWED"),
    ]
    
    for i, (label, value) in enumerate(ship_data):
        row = ship_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # 4. Goods Description
    add_section_heading("4. GOODS DESCRIPTION")
    doc.add_paragraph(lc.goods_description or "")
    
    doc.add_paragraph()  # Spacer
    
    # 5. Payment Terms
    add_section_heading("5. PAYMENT TERMS")
    
    payment_text = lc.payment_terms.value.upper()
    if lc.payment_terms.value == "usance" and lc.usance_days:
        payment_text = f"{lc.usance_days} DAYS FROM {lc.usance_from or 'B/L DATE'}"
    
    pay_table = doc.add_table(rows=5, cols=2)
    pay_table.style = 'Table Grid'
    
    pay_data = [
        ("Payment:", payment_text),
        ("Expiry Date:", lc.expiry_date.strftime("%d %B %Y") if lc.expiry_date else ""),
        ("Expiry Place:", lc.expiry_place or ""),
        ("Presentation Period:", f"{lc.presentation_period} days after shipment"),
        ("Confirmation:", lc.confirmation_instructions.value.upper()),
    ]
    
    for i, (label, value) in enumerate(pay_data):
        row = pay_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spacer
    
    # 6. Documents Required
    add_section_heading("6. DOCUMENTS REQUIRED")
    
    if lc.documents_required:
        for i, doc_item in enumerate(lc.documents_required, 1):
            if isinstance(doc_item, dict):
                doc_text = f"{i}. {doc_item.get('document_type', '')} - {doc_item.get('copies_original', 1)} original(s)"
                para = doc.add_paragraph(doc_text)
                if doc_item.get('specific_requirements'):
                    doc.add_paragraph(f"   {doc_item['specific_requirements']}")
            else:
                doc.add_paragraph(f"{i}. {doc_item}")
    
    doc.add_paragraph()  # Spacer
    
    # 7. Additional Conditions
    if lc.additional_conditions:
        add_section_heading("7. ADDITIONAL CONDITIONS")
        for cond in lc.additional_conditions:
            doc.add_paragraph(f"• {cond}")
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Generated on {datetime.now().strftime('%d %B %Y at %H:%M')} by TRDR Hub LC Builder"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.getvalue()
    
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=LC_{lc.reference_number}.docx"
        }
    )


# ============================================================================
# Bank-Specific PDF Export
# ============================================================================

BANK_FORMATS = {
    "SCB": {
        "name": "Standard Chartered Bank",
        "swift": "SCBLSGSG",
        "color": "#00A86B",  # SCB Green
        "header": "DOCUMENTARY CREDIT APPLICATION - STANDARD CHARTERED BANK FORMAT",
        "footer_text": "This document complies with Standard Chartered Bank LC submission requirements.",
        "requirements": [
            "All documents must be original or certified true copies",
            "Beneficiary name must match LC exactly (no abbreviations)",
            "Port names must include country codes",
        ],
    },
    "HSBC": {
        "name": "HSBC",
        "swift": "HSBCHKHH",
        "color": "#DB0011",  # HSBC Red
        "header": "DOCUMENTARY CREDIT APPLICATION - HSBC FORMAT",
        "footer_text": "This document complies with HSBC documentary credit submission standards.",
        "requirements": [
            "Documents must be in English or with certified English translation",
            "All amounts must be in figures and words",
            "Shipping date must be clearly visible",
        ],
    },
    "CITI": {
        "name": "Citibank",
        "swift": "CITIUS33",
        "color": "#003B70",  # Citi Blue
        "header": "DOCUMENTARY CREDIT APPLICATION - CITIBANK FORMAT",
        "footer_text": "This document complies with Citibank LC processing requirements.",
        "requirements": [
            "All documents must be dated and signed",
            "No alterations without authentication",
            "Copy documents must be marked COPY",
        ],
    },
    "DBS": {
        "name": "DBS Bank",
        "swift": "DBSSSGSG",
        "color": "#E31837",  # DBS Red
        "header": "DOCUMENTARY CREDIT APPLICATION - DBS BANK FORMAT",
        "footer_text": "This document complies with DBS Bank trade finance requirements.",
        "requirements": [
            "All amounts in LC currency only",
            "Clear description of goods required",
            "Transport documents must show shipped on board",
        ],
    },
}


@router.get("/bank-formats")
async def list_bank_formats():
    """List available bank-specific export formats"""
    return {
        "formats": [
            {
                "code": code,
                "name": fmt["name"],
                "swift": fmt["swift"],
                "requirements": fmt["requirements"],
            }
            for code, fmt in BANK_FORMATS.items()
        ]
    }


@router.post("/applications/{application_id}/export/pdf/{bank_code}")
async def export_bank_pdf(
    application_id: str,
    bank_code: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Export LC application as bank-specific formatted PDF"""
    from fastapi.responses import Response
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    
    if bank_code.upper() not in BANK_FORMATS:
        raise HTTPException(status_code=400, detail=f"Unknown bank format: {bank_code}")
    
    bank = BANK_FORMATS[bank_code.upper()]
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Generate bank-specific PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    bank_color = colors.HexColor(bank["color"])
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=10,
        alignment=1,
        textColor=bank_color
    )
    bank_header_style = ParagraphStyle(
        'BankHeader',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,
        textColor=colors.gray
    )
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=11,
        spaceBefore=12,
        spaceAfter=8,
        textColor=bank_color
    )
    normal_style = styles['Normal']
    
    elements = []
    
    # Bank Header
    elements.append(Paragraph(f"<b>{bank['name']}</b>", bank_header_style))
    elements.append(Paragraph(f"SWIFT: {bank['swift']}", bank_header_style))
    elements.append(Spacer(1, 10))
    
    # Title
    elements.append(Paragraph(bank["header"], title_style))
    elements.append(Paragraph(f"<b>Reference: {lc.reference_number}</b>", bank_header_style))
    elements.append(Spacer(1, 15))
    
    # Basic Info
    elements.append(Paragraph("1. CREDIT DETAILS", heading_style))
    basic_data = [
        ["Form of Credit:", lc.lc_type.value.upper()],
        ["Currency & Amount:", f"{lc.currency} {lc.amount:,.2f}"],
        ["Tolerance:", f"+{lc.tolerance_plus}% / -{lc.tolerance_minus}%"],
        ["Expiry Date:", lc.expiry_date.strftime("%d %B %Y") if lc.expiry_date else ""],
        ["Expiry Place:", lc.expiry_place or ""],
    ]
    basic_table = Table(basic_data, colWidths=[2*inch, 4*inch])
    basic_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f8f9fa")),
    ]))
    elements.append(basic_table)
    
    # Parties
    elements.append(Paragraph("2. PARTIES", heading_style))
    parties_data = [
        ["APPLICANT", "BENEFICIARY"],
        [lc.applicant_name or "", lc.beneficiary_name or ""],
        [lc.applicant_address or "", lc.beneficiary_address or ""],
        [lc.applicant_country or "", lc.beneficiary_country or ""],
    ]
    parties_table = Table(parties_data, colWidths=[3*inch, 3*inch])
    parties_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f8f9fa")),
    ]))
    elements.append(parties_table)
    
    # Shipment
    elements.append(Paragraph("3. SHIPMENT DETAILS", heading_style))
    ship_data = [
        ["Port of Loading:", lc.port_of_loading or "Any Port"],
        ["Port of Discharge:", lc.port_of_discharge or ""],
        ["Latest Shipment:", lc.latest_shipment_date.strftime("%d %B %Y") if lc.latest_shipment_date else ""],
        ["Incoterms:", f"{lc.incoterms} {lc.incoterms_place or ''}" if lc.incoterms else ""],
        ["Partial Shipments:", "ALLOWED" if lc.partial_shipments else "NOT ALLOWED"],
        ["Transhipment:", "ALLOWED" if lc.transhipment else "NOT ALLOWED"],
    ]
    ship_table = Table(ship_data, colWidths=[2*inch, 4*inch])
    ship_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f8f9fa")),
    ]))
    elements.append(ship_table)
    
    # Goods
    elements.append(Paragraph("4. DESCRIPTION OF GOODS/SERVICES", heading_style))
    goods_style = ParagraphStyle(
        'Goods',
        parent=normal_style,
        fontSize=9,
        borderColor=colors.lightgrey,
        borderWidth=0.5,
        borderPadding=8,
        backColor=colors.HexColor("#f8f9fa"),
    )
    elements.append(Paragraph(lc.goods_description or "", goods_style))
    
    # Payment
    elements.append(Paragraph("5. PAYMENT TERMS", heading_style))
    payment_text = lc.payment_terms.value.upper()
    if lc.payment_terms.value == "usance" and lc.usance_days:
        payment_text = f"{lc.usance_days} DAYS FROM {lc.usance_from or 'B/L DATE'}"
    pay_data = [
        ["Payment:", payment_text],
        ["Presentation Period:", f"{lc.presentation_period} days after shipment"],
        ["Confirmation:", lc.confirmation_instructions.value.upper()],
    ]
    pay_table = Table(pay_data, colWidths=[2*inch, 4*inch])
    pay_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f8f9fa")),
    ]))
    elements.append(pay_table)
    
    # Documents
    elements.append(Paragraph("6. DOCUMENTS REQUIRED", heading_style))
    if lc.documents_required:
        for i, doc_item in enumerate(lc.documents_required, 1):
            if isinstance(doc_item, dict):
                doc_text = f"{i}. {doc_item.get('document_type', '').replace('_', ' ').title()} - {doc_item.get('copies_original', 1)} original(s)"
                if doc_item.get('copies_copy'):
                    doc_text += f", {doc_item.get('copies_copy')} copy(ies)"
                elements.append(Paragraph(doc_text, normal_style))
            else:
                elements.append(Paragraph(f"{i}. {doc_item}", normal_style))
    
    # Additional Conditions
    if lc.additional_conditions:
        elements.append(Paragraph("7. ADDITIONAL CONDITIONS", heading_style))
        for cond in lc.additional_conditions:
            elements.append(Paragraph(f"• {cond}", normal_style))
    
    # Bank Requirements Box
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(f"<b>{bank['name']} SPECIFIC REQUIREMENTS</b>", heading_style))
    req_style = ParagraphStyle(
        'Requirements',
        parent=normal_style,
        fontSize=8,
        textColor=colors.HexColor("#6b7280"),
    )
    for req in bank["requirements"]:
        elements.append(Paragraph(f"• {req}", req_style))
    
    # Footer
    elements.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=normal_style,
        fontSize=8,
        textColor=colors.gray,
        alignment=1
    )
    elements.append(Paragraph(bank["footer_text"], footer_style))
    elements.append(Paragraph(
        f"Generated on {datetime.now().strftime('%d %B %Y at %H:%M')} by TRDR Hub LC Builder",
        footer_style
    ))
    
    doc.build(elements)
    
    buffer.seek(0)
    pdf_bytes = buffer.getvalue()
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=LC_{lc.reference_number}_{bank_code}.pdf"
        }
    )


# ============================================================================
# Email Notification Service for LC Builder
# ============================================================================

async def send_lc_status_notification(
    user_email: str,
    user_name: str,
    lc_reference: str,
    old_status: str,
    new_status: str,
    lc_amount: float,
    lc_currency: str,
    beneficiary_name: str,
    additional_info: Optional[str] = None
):
    """Send email notification when LC status changes."""
    from app.services.notifications import EmailService
    
    email_service = EmailService()
    
    status_messages = {
        "review": "Your LC application is now under review.",
        "submitted": "Your LC application has been submitted to the bank.",
        "approved": "Congratulations! Your LC application has been approved.",
        "rejected": "Your LC application requires attention.",
        "amended": "Your LC application has been amended.",
    }
    
    status_colors = {
        "draft": "#6B7280",
        "review": "#F59E0B",
        "submitted": "#3B82F6",
        "approved": "#10B981",
        "rejected": "#EF4444",
        "amended": "#8B5CF6",
    }
    
    subject = f"LC {lc_reference} - Status Update: {new_status.upper()}"
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 0; padding: 0; background: #f4f4f4; }}
            .container {{ max-width: 600px; margin: 20px auto; background: #fff; border-radius: 8px; overflow: hidden; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
            .header {{ background: linear-gradient(135deg, #10B981 0%, #059669 100%); color: white; padding: 30px; text-align: center; }}
            .header h1 {{ margin: 0; font-size: 24px; }}
            .content {{ padding: 30px; }}
            .status-badge {{ display: inline-block; padding: 8px 16px; border-radius: 20px; background: {status_colors.get(new_status, '#6B7280')}; color: white; font-weight: bold; font-size: 14px; text-transform: uppercase; }}
            .lc-details {{ background: #f9fafb; border-radius: 8px; padding: 20px; margin: 20px 0; }}
            .detail-row {{ display: flex; justify-content: space-between; padding: 8px 0; border-bottom: 1px solid #e5e7eb; }}
            .detail-row:last-child {{ border-bottom: none; }}
            .detail-label {{ color: #6B7280; }}
            .detail-value {{ font-weight: 600; color: #111827; }}
            .amount {{ font-size: 24px; color: #10B981; font-weight: bold; }}
            .cta {{ text-align: center; margin: 30px 0; }}
            .cta a {{ display: inline-block; padding: 12px 30px; background: #10B981; color: white; text-decoration: none; border-radius: 6px; font-weight: 600; }}
            .footer {{ background: #f9fafb; padding: 20px; text-align: center; color: #6B7280; font-size: 12px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>LC Status Update</h1>
            </div>
            <div class="content">
                <p>Hello {user_name},</p>
                <p>{status_messages.get(new_status, f'Your LC status has changed to {new_status}.')}</p>
                
                <div style="text-align: center; margin: 20px 0;">
                    <span class="status-badge">{new_status}</span>
                </div>
                
                <div class="lc-details">
                    <div class="detail-row">
                        <span class="detail-label">Reference Number</span>
                        <span class="detail-value">{lc_reference}</span>
                    </div>
                    <div class="detail-row">
                        <span class="detail-label">Amount</span>
                        <span class="detail-value amount">{lc_currency} {lc_amount:,.2f}</span>
                    </div>
                    <div class="detail-row">
                        <span class="detail-label">Beneficiary</span>
                        <span class="detail-value">{beneficiary_name}</span>
                    </div>
                    <div class="detail-row">
                        <span class="detail-label">Previous Status</span>
                        <span class="detail-value">{old_status.upper()}</span>
                    </div>
                </div>
                
                {f'<p style="background: #FEF3C7; padding: 12px; border-radius: 6px; border-left: 4px solid #F59E0B;">{additional_info}</p>' if additional_info else ''}
                
                <div class="cta">
                    <a href="https://trdrhub.com/lc-builder/dashboard">View LC Application</a>
                </div>
            </div>
            <div class="footer">
                <p>This is an automated notification from TRDR Hub LC Builder.</p>
                <p>© 2024 TRDR Hub. All rights reserved.</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    text = f"""
LC Status Update - {lc_reference}

Hello {user_name},

{status_messages.get(new_status, f'Your LC status has changed to {new_status}.')}

Status: {new_status.upper()}

LC Details:
- Reference: {lc_reference}
- Amount: {lc_currency} {lc_amount:,.2f}
- Beneficiary: {beneficiary_name}
- Previous Status: {old_status.upper()}

{additional_info or ''}

View your LC at: https://trdrhub.com/lc-builder/dashboard

---
TRDR Hub LC Builder
    """
    
    result = await email_service.send(
        to=user_email,
        subject=subject,
        html=html,
        text=text
    )
    
    return result


# ============================================================================
# Approval Workflow Endpoints
# ============================================================================

class StatusUpdateRequest(BaseModel):
    """Request to update LC status"""
    new_status: str = Field(..., description="New status: review, submitted, approved, rejected")
    comment: Optional[str] = None
    notify: bool = True


@router.post("/applications/{application_id}/status")
async def update_application_status(
    application_id: str,
    request: StatusUpdateRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Update the status of an LC application with optional notification.
    
    Status flow:
    - draft -> review -> submitted -> approved/rejected
    - approved/submitted -> amended (via amendment)
    """
    from fastapi import BackgroundTasks
    
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    old_status = lc.status.value
    new_status = request.new_status.lower()
    
    # Validate status transitions
    valid_transitions = {
        "draft": ["review", "submitted"],
        "review": ["draft", "submitted", "approved", "rejected"],
        "submitted": ["approved", "rejected", "amended"],
        "approved": ["amended"],
        "rejected": ["draft", "review"],
        "amended": ["review", "submitted"],
    }
    
    if new_status not in valid_transitions.get(old_status, []):
        raise HTTPException(
            status_code=400, 
            detail=f"Invalid status transition from {old_status} to {new_status}"
        )
    
    # Update status
    try:
        lc.status = LCStatus(new_status)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid status: {new_status}")
    
    lc.updated_at = datetime.utcnow()
    
    if new_status == "submitted":
        lc.submitted_at = datetime.utcnow()
    
    db.commit()
    
    # Send notification if enabled
    if request.notify:
        # Get user's email
        user = db.query(User).filter(User.id == current_user.id).first()
        if user and user.email:
            background_tasks.add_task(
                send_lc_status_notification,
                user_email=user.email,
                user_name=user.full_name or user.email,
                lc_reference=lc.reference_number,
                old_status=old_status,
                new_status=new_status,
                lc_amount=lc.amount,
                lc_currency=lc.currency,
                beneficiary_name=lc.beneficiary_name,
                additional_info=request.comment
            )
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "old_status": old_status,
        "new_status": new_status,
        "notification_sent": request.notify,
        "updated_at": lc.updated_at.isoformat()
    }


@router.post("/applications/{application_id}/submit-for-review")
async def submit_for_review(
    application_id: str,
    background_tasks: BackgroundTasks,
    reviewer_email: Optional[str] = None,
    notes: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Submit an LC application for review.
    Optionally send to a specific reviewer.
    """
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    if lc.status not in [LCStatus.DRAFT, LCStatus.AMENDED]:
        raise HTTPException(status_code=400, detail="Only draft or amended applications can be submitted for review")
    
    # Validate before submitting
    validation = validate_lc_application(lc)
    
    if not validation.is_valid:
        error_count = len([i for i in validation.issues if i.severity == "error"])
        if error_count > 0:
            raise HTTPException(
                status_code=400, 
                detail=f"Application has {error_count} error(s) that must be resolved before submission"
            )
    
    old_status = lc.status.value
    lc.status = LCStatus.REVIEW
    lc.updated_at = datetime.utcnow()
    db.commit()
    
    # Send notification to owner
    user = db.query(User).filter(User.id == current_user.id).first()
    if user and user.email:
        background_tasks.add_task(
            send_lc_status_notification,
            user_email=user.email,
            user_name=user.full_name or user.email,
            lc_reference=lc.reference_number,
            old_status=old_status,
            new_status="review",
            lc_amount=lc.amount,
            lc_currency=lc.currency,
            beneficiary_name=lc.beneficiary_name,
            additional_info=notes
        )
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "status": "review",
        "validation": {
            "is_valid": validation.is_valid,
            "risk_score": validation.risk_score,
            "issue_count": len(validation.issues),
        },
        "message": "Application submitted for review"
    }


@router.post("/applications/{application_id}/approve")
async def approve_application(
    application_id: str,
    background_tasks: BackgroundTasks,
    approval_notes: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Approve an LC application (for reviewers/managers).
    """
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    if lc.status not in [LCStatus.REVIEW, LCStatus.SUBMITTED]:
        raise HTTPException(status_code=400, detail="Only applications under review or submitted can be approved")
    
    old_status = lc.status.value
    lc.status = LCStatus.APPROVED
    lc.updated_at = datetime.utcnow()
    db.commit()
    
    # Send notification
    user = db.query(User).filter(User.id == current_user.id).first()
    if user and user.email:
        background_tasks.add_task(
            send_lc_status_notification,
            user_email=user.email,
            user_name=user.full_name or user.email,
            lc_reference=lc.reference_number,
            old_status=old_status,
            new_status="approved",
            lc_amount=lc.amount,
            lc_currency=lc.currency,
            beneficiary_name=lc.beneficiary_name,
            additional_info=approval_notes
        )
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "status": "approved",
        "approved_at": lc.updated_at.isoformat(),
        "message": "Application approved successfully"
    }


@router.post("/applications/{application_id}/reject")
async def reject_application(
    application_id: str,
    background_tasks: BackgroundTasks,
    rejection_reason: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Reject an LC application with reason.
    """
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    if lc.status not in [LCStatus.REVIEW, LCStatus.SUBMITTED]:
        raise HTTPException(status_code=400, detail="Only applications under review or submitted can be rejected")
    
    old_status = lc.status.value
    lc.status = LCStatus.REJECTED
    lc.updated_at = datetime.utcnow()
    db.commit()
    
    # Send notification with rejection reason
    user = db.query(User).filter(User.id == current_user.id).first()
    if user and user.email:
        background_tasks.add_task(
            send_lc_status_notification,
            user_email=user.email,
            user_name=user.full_name or user.email,
            lc_reference=lc.reference_number,
            old_status=old_status,
            new_status="rejected",
            lc_amount=lc.amount,
            lc_currency=lc.currency,
            beneficiary_name=lc.beneficiary_name,
            additional_info=f"Rejection Reason: {rejection_reason}"
        )
    
    return {
        "id": str(lc.id),
        "reference_number": lc.reference_number,
        "status": "rejected",
        "rejection_reason": rejection_reason,
        "message": "Application rejected"
    }


@router.get("/applications/{application_id}/workflow-status")
async def get_workflow_status(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get the current workflow status and available actions for an LC application.
    """
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    current_status = lc.status.value
    
    # Define available actions based on status
    status_actions = {
        "draft": [
            {"action": "submit_for_review", "label": "Submit for Review", "endpoint": "/submit-for-review"},
            {"action": "edit", "label": "Edit Application", "endpoint": None},
            {"action": "delete", "label": "Delete", "endpoint": None},
        ],
        "review": [
            {"action": "approve", "label": "Approve", "endpoint": "/approve"},
            {"action": "reject", "label": "Reject", "endpoint": "/reject"},
            {"action": "edit", "label": "Edit Application", "endpoint": None},
        ],
        "submitted": [
            {"action": "approve", "label": "Mark as Approved", "endpoint": "/approve"},
            {"action": "reject", "label": "Mark as Rejected", "endpoint": "/reject"},
        ],
        "approved": [
            {"action": "amend", "label": "Create Amendment", "endpoint": "/amend"},
            {"action": "export", "label": "Export Documents", "endpoint": None},
        ],
        "rejected": [
            {"action": "resubmit", "label": "Resubmit", "endpoint": "/submit-for-review"},
            {"action": "edit", "label": "Edit and Fix", "endpoint": None},
        ],
        "amended": [
            {"action": "submit_for_review", "label": "Submit Amendment for Review", "endpoint": "/submit-for-review"},
            {"action": "edit", "label": "Edit Amendment", "endpoint": None},
        ],
    }
    
    # Workflow timeline
    workflow_steps = [
        {"step": "draft", "label": "Draft", "completed": current_status != "draft"},
        {"step": "review", "label": "Under Review", "completed": current_status in ["submitted", "approved", "rejected"]},
        {"step": "submitted", "label": "Submitted to Bank", "completed": current_status in ["approved", "rejected"]},
        {"step": "approved", "label": "Approved", "completed": current_status == "approved"},
    ]
    
    return {
        "application_id": str(lc.id),
        "reference_number": lc.reference_number,
        "current_status": current_status,
        "available_actions": status_actions.get(current_status, []),
        "workflow_steps": workflow_steps,
        "created_at": lc.created_at.isoformat(),
        "updated_at": lc.updated_at.isoformat(),
        "submitted_at": lc.submitted_at.isoformat() if lc.submitted_at else None,
    }


# ============================================================================
# Team Collaboration - Sharing
# ============================================================================

from app.models.lc_builder import LCShare, SharePermission, LCComment, NotificationPreference


class ShareRequest(BaseModel):
    """Request to share an LC application"""
    email: str = Field(..., description="Email of person to share with")
    permission: str = Field(default="view", description="Permission level: view, comment, edit, review")
    message: Optional[str] = None


@router.post("/applications/{application_id}/share")
async def share_application(
    application_id: str,
    request: ShareRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Share an LC application with a colleague."""
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    # Check if already shared with this email
    existing = db.query(LCShare).filter(
        LCShare.lc_application_id == application_id,
        LCShare.shared_with_email == request.email.lower()
    ).first()
    
    if existing:
        # Update permission
        existing.permission = SharePermission(request.permission)
        db.commit()
        return {"status": "updated", "share_id": str(existing.id)}
    
    # Check if email belongs to an existing user
    shared_user = db.query(User).filter(User.email == request.email.lower()).first()
    
    # Create share
    share = LCShare(
        lc_application_id=lc.id,
        shared_by_id=current_user.id,
        shared_with_user_id=shared_user.id if shared_user else None,
        shared_with_email=request.email.lower(),
        permission=SharePermission(request.permission),
        is_accepted=True if shared_user else False,
    )
    
    db.add(share)
    db.commit()
    db.refresh(share)
    
    # Send notification email
    from app.services.notifications import EmailService
    email_service = EmailService()
    
    share_html = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <h2 style="color: #10B981;">LC Application Shared With You</h2>
        <p>{current_user.full_name or current_user.email} has shared an LC application with you.</p>
        <div style="background: #f3f4f6; padding: 20px; border-radius: 8px; margin: 20px 0;">
            <p><strong>Reference:</strong> {lc.reference_number}</p>
            <p><strong>Amount:</strong> {lc.currency} {lc.amount:,.2f}</p>
            <p><strong>Beneficiary:</strong> {lc.beneficiary_name}</p>
            <p><strong>Permission:</strong> {request.permission.upper()}</p>
        </div>
        {f'<p><em>Message: {request.message}</em></p>' if request.message else ''}
        <a href="https://trdrhub.com/lc-builder/dashboard" 
           style="display: inline-block; background: #10B981; color: white; padding: 12px 24px; 
                  text-decoration: none; border-radius: 6px; margin-top: 20px;">
            View LC Application
        </a>
    </div>
    """
    
    background_tasks.add_task(
        email_service.send,
        to=request.email,
        subject=f"LC Application Shared: {lc.reference_number}",
        html=share_html,
        text=f"{current_user.full_name or current_user.email} shared LC {lc.reference_number} with you."
    )
    
    return {
        "status": "shared",
        "share_id": str(share.id),
        "email": request.email,
        "permission": request.permission,
        "is_existing_user": shared_user is not None
    }


@router.get("/applications/{application_id}/shares")
async def list_shares(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all shares for an LC application."""
    # Check access (owner or shared with)
    lc = db.query(LCApplication).filter(LCApplication.id == application_id).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    is_owner = str(lc.user_id) == str(current_user.id)
    is_shared = db.query(LCShare).filter(
        LCShare.lc_application_id == application_id,
        LCShare.shared_with_user_id == current_user.id
    ).first() is not None
    
    if not is_owner and not is_shared:
        raise HTTPException(status_code=403, detail="Access denied")
    
    shares = db.query(LCShare).filter(LCShare.lc_application_id == application_id).all()
    
    result = []
    for s in shares:
        # Get user info if available
        shared_user = None
        if s.shared_with_user_id:
            shared_user = db.query(User).filter(User.id == s.shared_with_user_id).first()
        
        result.append({
            "id": str(s.id),
            "email": s.shared_with_email,
            "user_name": shared_user.full_name if shared_user else None,
            "permission": s.permission.value,
            "is_accepted": s.is_accepted,
            "invited_at": s.invited_at.isoformat() if s.invited_at else None,
            "accepted_at": s.accepted_at.isoformat() if s.accepted_at else None,
        })
    
    return {"shares": result, "is_owner": is_owner}


@router.delete("/applications/{application_id}/shares/{share_id}")
async def remove_share(
    application_id: str,
    share_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Remove a share from an LC application."""
    lc = db.query(LCApplication).filter(
        LCApplication.id == application_id,
        LCApplication.user_id == current_user.id
    ).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found or access denied")
    
    share = db.query(LCShare).filter(
        LCShare.id == share_id,
        LCShare.lc_application_id == application_id
    ).first()
    
    if not share:
        raise HTTPException(status_code=404, detail="Share not found")
    
    db.delete(share)
    db.commit()
    
    return {"status": "removed"}


@router.get("/shared-with-me")
async def list_shared_with_me(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all LC applications shared with the current user."""
    shares = db.query(LCShare).filter(
        (LCShare.shared_with_user_id == current_user.id) |
        (LCShare.shared_with_email == current_user.email)
    ).all()
    
    result = []
    for s in shares:
        lc = db.query(LCApplication).filter(LCApplication.id == s.lc_application_id).first()
        if lc:
            owner = db.query(User).filter(User.id == lc.user_id).first()
            result.append({
                "share_id": str(s.id),
                "lc_id": str(lc.id),
                "reference_number": lc.reference_number,
                "name": lc.name,
                "status": lc.status.value,
                "amount": lc.amount,
                "currency": lc.currency,
                "beneficiary_name": lc.beneficiary_name,
                "permission": s.permission.value,
                "shared_by": owner.full_name if owner else "Unknown",
                "shared_at": s.invited_at.isoformat() if s.invited_at else None,
            })
    
    return {"shared_applications": result, "count": len(result)}


# ============================================================================
# Team Collaboration - Comments
# ============================================================================

class CommentRequest(BaseModel):
    """Request to add a comment"""
    content: str = Field(..., min_length=1, max_length=5000)
    parent_id: Optional[str] = None
    field_reference: Optional[str] = None
    mentions: List[str] = Field(default_factory=list)


@router.post("/applications/{application_id}/comments")
async def add_comment(
    application_id: str,
    request: CommentRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Add a comment to an LC application."""
    # Check access
    lc = db.query(LCApplication).filter(LCApplication.id == application_id).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    is_owner = str(lc.user_id) == str(current_user.id)
    share = db.query(LCShare).filter(
        LCShare.lc_application_id == application_id,
        LCShare.shared_with_user_id == current_user.id,
        LCShare.permission.in_([SharePermission.COMMENT, SharePermission.EDIT, SharePermission.REVIEW, SharePermission.ADMIN])
    ).first()
    
    if not is_owner and not share:
        raise HTTPException(status_code=403, detail="You don't have permission to comment")
    
    comment = LCComment(
        lc_application_id=lc.id,
        user_id=current_user.id,
        content=request.content,
        parent_id=request.parent_id,
        field_reference=request.field_reference,
        mentions=request.mentions,
    )
    
    db.add(comment)
    db.commit()
    db.refresh(comment)
    
    # Notify mentioned users and owner
    from app.services.notifications import EmailService
    email_service = EmailService()
    
    # Notify owner if commenter is not owner
    if not is_owner:
        owner = db.query(User).filter(User.id == lc.user_id).first()
        if owner and owner.email:
            background_tasks.add_task(
                email_service.send,
                to=owner.email,
                subject=f"New Comment on LC {lc.reference_number}",
                html=f"""
                <p>{current_user.full_name or current_user.email} commented on your LC application.</p>
                <blockquote style="background: #f3f4f6; padding: 12px; border-left: 4px solid #10B981;">
                    {request.content}
                </blockquote>
                <a href="https://trdrhub.com/lc-builder/dashboard/edit/{lc.id}">View Application</a>
                """,
                text=f"New comment from {current_user.full_name}: {request.content}"
            )
    
    return {
        "id": str(comment.id),
        "content": comment.content,
        "user_id": str(comment.user_id),
        "user_name": current_user.full_name or current_user.email,
        "created_at": comment.created_at.isoformat(),
        "field_reference": comment.field_reference,
    }


@router.get("/applications/{application_id}/comments")
async def list_comments(
    application_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all comments on an LC application."""
    # Check access
    lc = db.query(LCApplication).filter(LCApplication.id == application_id).first()
    
    if not lc:
        raise HTTPException(status_code=404, detail="LC application not found")
    
    is_owner = str(lc.user_id) == str(current_user.id)
    is_shared = db.query(LCShare).filter(
        LCShare.lc_application_id == application_id,
        LCShare.shared_with_user_id == current_user.id
    ).first() is not None
    
    if not is_owner and not is_shared:
        raise HTTPException(status_code=403, detail="Access denied")
    
    comments = db.query(LCComment).filter(
        LCComment.lc_application_id == application_id
    ).order_by(LCComment.created_at.desc()).all()
    
    result = []
    for c in comments:
        user = db.query(User).filter(User.id == c.user_id).first()
        result.append({
            "id": str(c.id),
            "content": c.content,
            "user_id": str(c.user_id),
            "user_name": user.full_name if user else "Unknown",
            "user_email": user.email if user else None,
            "parent_id": str(c.parent_id) if c.parent_id else None,
            "field_reference": c.field_reference,
            "mentions": c.mentions,
            "is_resolved": c.is_resolved,
            "created_at": c.created_at.isoformat(),
        })
    
    return {"comments": result, "count": len(result)}


@router.post("/applications/{application_id}/comments/{comment_id}/resolve")
async def resolve_comment(
    application_id: str,
    comment_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Mark a comment as resolved."""
    comment = db.query(LCComment).filter(
        LCComment.id == comment_id,
        LCComment.lc_application_id == application_id
    ).first()
    
    if not comment:
        raise HTTPException(status_code=404, detail="Comment not found")
    
    comment.is_resolved = True
    comment.resolved_by = current_user.id
    comment.resolved_at = datetime.utcnow()
    db.commit()
    
    return {"status": "resolved"}


# ============================================================================
# Notification Preferences
# ============================================================================

class NotificationPreferencesUpdate(BaseModel):
    """Update notification preferences"""
    email_on_status_change: Optional[bool] = None
    email_on_share: Optional[bool] = None
    email_on_comment: Optional[bool] = None
    email_on_mention: Optional[bool] = None
    email_on_approval_request: Optional[bool] = None
    email_on_rejection: Optional[bool] = None
    email_digest: Optional[str] = None  # instant, daily, weekly, none
    in_app_on_status_change: Optional[bool] = None
    in_app_on_share: Optional[bool] = None
    in_app_on_comment: Optional[bool] = None
    in_app_on_mention: Optional[bool] = None


@router.get("/notification-preferences")
async def get_notification_preferences(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get current user's notification preferences."""
    prefs = db.query(NotificationPreference).filter(
        NotificationPreference.user_id == current_user.id
    ).first()
    
    if not prefs:
        # Return defaults
        return {
            "email_on_status_change": True,
            "email_on_share": True,
            "email_on_comment": True,
            "email_on_mention": True,
            "email_on_approval_request": True,
            "email_on_rejection": True,
            "email_digest": "instant",
            "in_app_on_status_change": True,
            "in_app_on_share": True,
            "in_app_on_comment": True,
            "in_app_on_mention": True,
        }
    
    return {
        "email_on_status_change": prefs.email_on_status_change,
        "email_on_share": prefs.email_on_share,
        "email_on_comment": prefs.email_on_comment,
        "email_on_mention": prefs.email_on_mention,
        "email_on_approval_request": prefs.email_on_approval_request,
        "email_on_rejection": prefs.email_on_rejection,
        "email_digest": prefs.email_digest,
        "in_app_on_status_change": prefs.in_app_on_status_change,
        "in_app_on_share": prefs.in_app_on_share,
        "in_app_on_comment": prefs.in_app_on_comment,
        "in_app_on_mention": prefs.in_app_on_mention,
    }


@router.put("/notification-preferences")
async def update_notification_preferences(
    request: NotificationPreferencesUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update notification preferences."""
    prefs = db.query(NotificationPreference).filter(
        NotificationPreference.user_id == current_user.id
    ).first()
    
    if not prefs:
        prefs = NotificationPreference(user_id=current_user.id)
        db.add(prefs)
    
    # Update only provided fields
    for field, value in request.dict(exclude_unset=True).items():
        if value is not None:
            setattr(prefs, field, value)
    
    db.commit()
    db.refresh(prefs)
    
    return {"status": "updated"}

